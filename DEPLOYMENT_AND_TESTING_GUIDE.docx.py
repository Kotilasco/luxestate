"""Generate ZAI-CTS Deployment and Testing Guide DOCX."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code_text):
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def add_test_row(table, step, action, expected, status=""):
    """Add a test case row to a table."""
    row = table.add_row().cells
    row[0].text = str(step)
    row[1].text = action
    row[2].text = expected
    row[3].text = status
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    return row

doc = Document()

# Title
title = doc.add_heading("ZAI-CTS Deployment & Testing Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0d, 0x94, 0x88)

subtitle = doc.add_paragraph("How to run the Zimbabwe AI-Enhanced Carbon Trading Ecosystem on another computer")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph(f"Generated: 2026-06-26  |  Repository: https://github.com/Kotilasco/luxestate")
doc.add_paragraph()

# =============================================================================
# 1. SYSTEM OVERVIEW
# =============================================================================
doc.add_heading("1. System Overview", level=1)
doc.add_paragraph(
    "ZAI-CTS is a production-grade national carbon trading platform built on a microservices architecture. "
    "It consists of a Next.js frontend, a Node.js API Gateway, multiple Python FastAPI backend services, "
    "a PostgreSQL/PostGIS database, and optional Hyperledger Fabric blockchain components."
)

arch_table = doc.add_table(rows=1, cols=4)
arch_table.style = "Table Grid"
headers = ["Component", "Technology", "Port", "Purpose"]
for i, h in enumerate(headers):
    arch_table.rows[0].cells[i].text = h
    arch_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(arch_table.rows[0].cells[i], "0f766e")
    arch_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

components = [
    ("Web Portal", "Next.js 15 + TypeScript + MUI v6", "3000", "User interface for all registry operations"),
    ("API Gateway", "Fastify + Node.js", "8082", "Routes requests to backend microservices"),
    ("Carbon Registry", "FastAPI + SQLAlchemy", "8102", "Core carbon project & credit registry"),
    ("AI Validation", "FastAPI + LangChain", "8103", "AI-powered project validation & reports"),
    ("GIS Service", "FastAPI + PostGIS", "8104", "Spatial analysis & boundary validation"),
    ("Marketplace", "FastAPI", "8106", "Carbon credit trading & settlement"),
    ("Compliance", "FastAPI", "8107", "Regulatory compliance & enforcement"),
    ("PostgreSQL", "PostgreSQL 17 + PostGIS", "5432", "Primary database with spatial extensions"),
    ("Docker", "Docker + Compose", "N/A", "Container orchestration for local development"),
]
for comp, tech, port, purpose in components:
    row = arch_table.add_row().cells
    row[0].text = comp
    row[1].text = tech
    row[2].text = port
    row[3].text = purpose
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# =============================================================================
# 2. PREREQUISITES
# =============================================================================
doc.add_heading("2. Prerequisites", level=1)
doc.add_paragraph("Before starting, ensure the following software is installed on the target machine:")

prereq = [
    ("Git", "2.40+", "https://git-scm.com/downloads", "Clone repository"),
    ("Node.js", "20.x LTS", "https://nodejs.org/", "Frontend + API Gateway"),
    ("Python", "3.11+", "https://www.python.org/downloads/", "Backend microservices"),
    ("PostgreSQL", "15+", "https://www.postgresql.org/download/", "Primary database"),
    ("PostGIS", "3.4+", "https://postgis.net/install/", "Spatial extensions"),
    ("Docker", "28.x", "https://www.docker.com/products/docker-desktop/", "Container runtime (optional)"),
    ("Docker Compose", "2.38+", "Included with Docker Desktop", "Multi-container orchestration"),
]

prereq_table = doc.add_table(rows=1, cols=4)
prereq_table.style = "Table Grid"
headers = ["Software", "Version", "Download", "Required For"]
for i, h in enumerate(headers):
    prereq_table.rows[0].cells[i].text = h
    prereq_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(prereq_table.rows[0].cells[i], "0f766e")
    prereq_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for sw, ver, dl, req in prereq:
    row = prereq_table.add_row().cells
    row[0].text = sw
    row[1].text = ver
    row[2].text = dl
    row[3].text = req
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph("Verify installations:")
add_code_block(doc, "git --version\nnode --version\npython --version\npsql --version\ndocker --version\ndocker compose version")

doc.add_page_break()

# =============================================================================
# 3. CLONE & INITIAL SETUP
# =============================================================================
doc.add_heading("3. Clone & Initial Setup", level=1)

doc.add_heading("3.1 Clone the Repository", level=2)
add_code_block(doc, "git clone https://github.com/Kotilasco/luxestate.git\ncd luxestate")

doc.add_heading("3.2 Create PostgreSQL Database", level=2)
doc.add_paragraph("Create the database and enable required extensions:")
add_code_block(doc, """-- As postgres superuser or via psql:
CREATE DATABASE zai_cts;
\\c zai_cts
CREATE EXTENSION IF NOT EXISTS \"uuid-ossp\";
CREATE EXTENSION IF NOT EXISTS postgis;
CREATE SCHEMA IF NOT EXISTS identity;
CREATE SCHEMA IF NOT EXISTS registry;
CREATE SCHEMA IF NOT EXISTS gis;
CREATE SCHEMA IF NOT EXISTS audit;
CREATE SCHEMA IF NOT EXISTS market;
CREATE SCHEMA IF NOT EXISTS compliance;
CREATE SCHEMA IF NOT EXISTS article6;
CREATE SCHEMA IF NOT EXISTS reporting;""")

doc.add_heading("3.3 Run Database Migrations", level=2)
doc.add_paragraph("Apply the core schema and enterprise domains:")
add_code_block(doc, "psql -U your_username -d zai_cts -f database/schemas/core.sql\npsql -U your_username -d zai_cts -f database/migrations/0001_initial_core.sql\npsql -U your_username -d zai_cts -f database/migrations/0002_enterprise_domains.sql")

doc.add_paragraph("Note: If running on Windows with Git Bash, use double slashes for paths:")
add_code_block(doc, "psql -U postgres -d zai_cts -f database/schemas/core.sql")

doc.add_heading("3.4 Environment Configuration", level=2)
doc.add_paragraph("Create a .env file in the project root with database credentials:")
add_code_block(doc, """DATABASE_URL=postgresql+asyncpg://postgres:your_password@localhost:5432/zai_cts
POSTGRES_USER=postgres
POSTGRES_PASSWORD=your_password
POSTGRES_DB=zai_cts
POSTGRES_HOST=localhost
POSTGRES_PORT=5432""")

doc.add_paragraph("For the Carbon Registry service, create .env in backend/services/carbon-registry-service/:")
add_code_block(doc, """DATABASE_URL=postgresql+asyncpg://postgres:your_password@localhost:5432/zai_cts
SECRET_KEY=your-super-secret-key-min-32-characters-long
LOG_LEVEL=INFO""")

doc.add_page_break()

# =============================================================================
# 4. BACKEND SERVICES SETUP
# =============================================================================
doc.add_heading("4. Backend Services Setup", level=1)
doc.add_paragraph("Each backend service runs independently. Setup follows the same pattern for all.")

doc.add_heading("4.1 Carbon Registry Service (Port 8102)", level=2)
add_code_block(doc, """cd backend/services/carbon-registry-service
python -m venv .venv

# Windows:
.venv\\Scripts\\activate

# macOS/Linux:
source .venv/bin/activate

pip install -r requirements.txt
python -m uvicorn app.main:app --host 0.0.0.0 --port 8102 --reload""")

doc.add_heading("4.2 AI Validation Service (Port 8103)", level=2)
add_code_block(doc, """cd backend/services/ai-validation-service
python -m venv .venv
.venv\\Scripts\\activate  # or source .venv/bin/activate
pip install -r requirements.txt
python -m uvicorn app.main:app --host 0.0.0.0 --port 8103 --reload""")

doc.add_heading("4.3 GIS Service (Port 8104)", level=2)
add_code_block(doc, """cd backend/services/gis-service
python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt
python -m uvicorn app.main:app --host 0.0.0.0 --port 8104 --reload""")

doc.add_heading("4.4 Marketplace Service (Port 8106)", level=2)
add_code_block(doc, """cd backend/services/marketplace-service
python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt
python -m uvicorn app.main:app --host 0.0.0.0 --port 8106 --reload""")

doc.add_heading("4.5 Compliance Service (Port 8107)", level=2)
add_code_block(doc, """cd backend/services/compliance-service
python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt
python -m uvicorn app.main:app --host 0.0.0.0 --port 8107 --reload""")

doc.add_paragraph()
doc.add_paragraph("Alternative: Use the provided batch script (Windows only):")
add_code_block(doc, "START_SERVERS.bat")

doc.add_page_break()

# =============================================================================
# 5. API GATEWAY SETUP
# =============================================================================
doc.add_heading("5. API Gateway Setup", level=1)
doc.add_paragraph("The API Gateway routes frontend requests to the correct backend service.")

add_code_block(doc, """cd api-gateway
npm install
npm run dev""")

doc.add_paragraph("The gateway will start on port 8082. Verify it's running:")
add_code_block(doc, "curl http://localhost:8082/health")

doc.add_paragraph("Expected response:")
add_code_block(doc, '{"status":"healthy","service":"zai-cts-api-gateway"}')

doc.add_heading("5.1 Gateway Route Mapping", level=2)
route_table = doc.add_table(rows=1, cols=3)
route_table.style = "Table Grid"
headers = ["Frontend Path", "Backend Service", "Port"]
for i, h in enumerate(headers):
    route_table.rows[0].cells[i].text = h
    route_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(route_table.rows[0].cells[i], "0f766e")
    route_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

routes = [
    ("/api/v1/auth/*", "Carbon Registry", "8102"),
    ("/api/v1/projects", "Carbon Registry", "8102"),
    ("/api/v1/national-*", "Carbon Registry", "8102"),
    ("/anchors/*", "Carbon Registry", "8102"),
    ("/api/v1/gis/*", "GIS Service", "8104"),
    ("/api/v1/marketplace/*", "Marketplace", "8106"),
    ("/api/v1/compliance/*", "Compliance", "8107"),
    ("/api/v1/ai/*", "AI Validation", "8103"),
]
for path, svc, port in routes:
    row = route_table.add_row().cells
    row[0].text = path
    row[1].text = svc
    row[2].text = port
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# =============================================================================
# 6. FRONTEND SETUP
# =============================================================================
doc.add_heading("6. Frontend Setup", level=1)
doc.add_paragraph("The web portal is a Next.js 15 application with TypeScript and Material-UI v6.")

add_code_block(doc, """cd frontend/web-portal
npm install
npm run dev""")

doc.add_paragraph("The frontend will start on port 3000 (or 3004 if 3000 is in use). Open:")
add_code_block(doc, "http://localhost:3000")

doc.add_heading("6.1 Frontend Environment", level=2)
doc.add_paragraph("Create .env.local in frontend/web-portal/:")
add_code_block(doc, """NEXT_PUBLIC_API_URL=http://localhost:8082
NEXT_PUBLIC_APP_NAME=ZAI-CTS""")

doc.add_heading("6.2 Known Issue: next/dynamic in Server Components", level=2)
doc.add_paragraph(
    "If you encounter the error: `ssr: false is not allowed with next/dynamic in Server Components`, "
    "the page.tsx uses dynamic imports with ssr: false. This is a known Next.js 15 constraint. "
    "The current workaround is that the app already includes a 'use client' directive and ClientOnly wrapper. "
    "If the error persists, ensure the RegistryConsole component is imported as a client component."
)

doc.add_page_break()

# =============================================================================
# 7. DOCKER COMPOSE (ALTERNATIVE)
# =============================================================================
doc.add_heading("7. Docker Compose (Alternative)", level=1)
doc.add_paragraph("For a fully containerized setup, use Docker Compose:")

add_code_block(doc, "docker compose -f infrastructure/docker/docker-compose.yml up --build")

doc.add_paragraph("This will start PostgreSQL, all backend services, the API Gateway, and the frontend in containers.")

doc.add_heading("7.1 Verify Container Status", level=2)
add_code_block(doc, "docker ps --format \"table {{.Names}}\\t{{.Status}}\\t{{.Ports}}\"")

doc.add_page_break()

# =============================================================================
# 8. STARTUP ORDER
# =============================================================================
doc.add_heading("8. Startup Order", level=1)
doc.add_paragraph("Services must be started in this order to avoid connection errors:")

order_table = doc.add_table(rows=1, cols=4)
order_table.style = "Table Grid"
headers = ["Order", "Service", "Command", "Wait For"]
for i, h in enumerate(headers):
    order_table.rows[0].cells[i].text = h
    order_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(order_table.rows[0].cells[i], "0f766e")
    order_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

startup_order = [
    ("1", "PostgreSQL", "docker start zetdc_bts-postgres-1  OR  pg_ctl start", "N/A"),
    ("2", "Carbon Registry", "uvicorn app.main:app --port 8102", "PostgreSQL"),
    ("3", "AI Validation", "uvicorn app.main:app --port 8103", "PostgreSQL"),
    ("4", "GIS Service", "uvicorn app.main:app --port 8104", "PostgreSQL"),
    ("5", "Marketplace", "uvicorn app.main:app --port 8106", "PostgreSQL"),
    ("6", "Compliance", "uvicorn app.main:app --port 8107", "PostgreSQL"),
    ("7", "API Gateway", "npm run dev", "All backends"),
    ("8", "Frontend", "npm run dev", "API Gateway"),
]
for order, svc, cmd, wait in startup_order:
    row = order_table.add_row().cells
    row[0].text = order
    row[1].text = svc
    row[2].text = cmd
    row[3].text = wait
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph("Quick verification after startup:")
add_code_block(doc, """curl http://localhost:8082/health          # Gateway
curl http://localhost:8102/health          # Carbon Registry
curl http://localhost:8103/health          # AI Validation
curl http://localhost:8104/health          # GIS
curl http://localhost:8106/health          # Marketplace
curl http://localhost:8107/health          # Compliance""")

doc.add_page_break()

# =============================================================================
# 9. END-TO-END TEST PROCEDURES
# =============================================================================
doc.add_heading("9. End-to-End Test Procedures", level=1)
doc.add_paragraph(
    "These tests verify the entire system works correctly from the user's perspective. "
    "Execute them in order. Record results in the Pass/Fail column."
)

doc.add_heading("9.1 Database Connectivity Tests", level=2)
db_table = doc.add_table(rows=1, cols=4)
db_table.style = "Table Grid"
headers = ["Step", "Action", "Expected Result", "Pass/Fail"]
for i, h in enumerate(headers):
    db_table.rows[0].cells[i].text = h
    db_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(db_table.rows[0].cells[i], "0f766e")
    db_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

db_tests = [
    ("Connect to PostgreSQL", "psql -U postgres -d zai_cts -c 'SELECT 1;'", "Returns 1 row with value 1"),
    ("Check schemas exist", "\\dn in psql", "identity, registry, gis, audit, market, compliance, article6, reporting listed"),
    ("Check PostGIS", "SELECT PostGIS_Version();", "Returns PostGIS version string"),
    ("Check core tables", "\\dt registry.* in psql", "carbon_projects, carbon_credit_batches, anchor_batches, carbon_credit_entries visible"),
    ("Check seed data", "SELECT * FROM identity.organizations LIMIT 1;", "Zimbabwe Climate Authority record exists"),
    ("Check users table", "SELECT count(*) FROM identity.users;", "Returns count >= 1 (seed admin)"),
]
for i, (action, cmd, expected) in enumerate(db_tests, 1):
    add_test_row(db_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.2 Backend API Health Tests", level=2)
health_table = doc.add_table(rows=1, cols=4)
health_table.style = "Table Grid"
for i, h in enumerate(headers):
    health_table.rows[0].cells[i].text = h
    health_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(health_table.rows[0].cells[i], "0f766e")
    health_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

health_tests = [
    ("Carbon Registry health", "curl http://localhost:8102/health", '{"status":"ok"} or similar healthy response'),
    ("AI Validation health", "curl http://localhost:8103/health", "Healthy response"),
    ("GIS Service health", "curl http://localhost:8104/health", "Healthy response"),
    ("Marketplace health", "curl http://localhost:8106/health", "Healthy response"),
    ("Compliance health", "curl http://localhost:8107/health", "Healthy response"),
    ("API Gateway health", "curl http://localhost:8082/health", '{"status":"healthy","service":"zai-cts-api-gateway"}'),
]
for i, (action, cmd, expected) in enumerate(health_tests, 1):
    add_test_row(health_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.3 Authentication Tests", level=2)
auth_table = doc.add_table(rows=1, cols=4)
auth_table.style = "Table Grid"
for i, h in enumerate(headers):
    auth_table.rows[0].cells[i].text = h
    auth_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(auth_table.rows[0].cells[i], "0f766e")
    auth_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

auth_tests = [
    ("Seed admin login", "POST /api/v1/auth/login with admin@zai-cts.gov.zw / ZAI-CTS-Admin-2026!", "Returns access_token, refresh_token, user object"),
    ("Get current user", "GET /api/v1/auth/me with Bearer token", "Returns user profile with role"),
    ("List all users", "GET /api/v1/auth/users", "Returns array of users (admin only)"),
    ("Invalid password", "POST /api/v1/auth/login with wrong password", "Returns 401 Unauthorized"),
    ("Missing credentials", "POST /api/v1/auth/login with empty body", "Returns 422 Validation Error"),
    ("Password reset request", "POST /api/v1/auth/forgot-password with valid email", "Returns success message"),
]
for i, (action, cmd, expected) in enumerate(auth_tests, 1):
    add_test_row(auth_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.4 Carbon Registry Tests", level=2)
registry_table = doc.add_table(rows=1, cols=4)
registry_table.style = "Table Grid"
for i, h in enumerate(headers):
    registry_table.rows[0].cells[i].text = h
    registry_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(registry_table.rows[0].cells[i], "0f766e")
    registry_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

registry_tests = [
    ("List projects", "GET /api/v1/projects", "Returns array of carbon projects"),
    ("Get project by ID", "GET /api/v1/projects/{project_id}", "Returns single project with all fields"),
    ("Register new project", "POST /api/v1/projects with valid payload", "Returns created project with UUID"),
    ("Project workflow advance", "POST /api/v1/projects/{id}/workflow", "Returns updated project status"),
    ("Issue credit batch", "POST /api/v1/projects/{id}/credits", "Returns batch with serial_prefix"),
    ("List credit batches", "GET /api/v1/projects/{id}/credits", "Returns array of credit batches"),
    ("Get national readiness", "GET /api/v1/national-readiness", "Returns maturity score and stages"),
    ("Get national operations", "GET /api/v1/national-operations", "Returns registry accounts, methodologies, etc."),
]
for i, (action, cmd, expected) in enumerate(registry_tests, 1):
    add_test_row(registry_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.5 Merkle Anchoring Tests", level=2)
anchor_table = doc.add_table(rows=1, cols=4)
anchor_table.style = "Table Grid"
for i, h in enumerate(headers):
    anchor_table.rows[0].cells[i].text = h
    anchor_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(anchor_table.rows[0].cells[i], "0f766e")
    anchor_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

anchor_tests = [
    ("Unanchored count", "GET /anchors/status/unanchored-count", "Returns {\"unanchored_count\": N}"),
    ("Create anchor batch (empty)", "POST /anchors/batch when count=0", "Returns 400 with \"no_entries\" status"),
    ("Create anchor batch (with data)", "POST /anchors/batch after creating credit entries", "Returns anchor_id, merkle_root, fabric_tx_id"),
    ("Verify anchor", "GET /anchors/{anchor_id}/verify", "Returns is_valid=true, stored and recomputed roots match"),
    ("Reconcile chain", "GET /anchors/reconcile", "Returns array of anchor verification results"),
    ("Gateway routes anchors", "curl http://localhost:8082/anchors/status/unanchored-count", "Same result as direct backend call"),
    ("Frontend anchoring tab", "Open browser, login, click Anchoring", "Shows unanchored count, create batch, verify, reconcile UI"),
]
for i, (action, cmd, expected) in enumerate(anchor_tests, 1):
    add_test_row(anchor_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.6 Frontend UI Tests", level=2)
ui_table = doc.add_table(rows=1, cols=4)
ui_table.style = "Table Grid"
for i, h in enumerate(headers):
    ui_table.rows[0].cells[i].text = h
    ui_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(ui_table.rows[0].cells[i], "0f766e")
    ui_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

ui_tests = [
    ("Page loads", "Open http://localhost:3000", "Login page loads without errors"),
    ("Login success", "Enter admin@zai-cts.gov.zw / ZAI-CTS-Admin-2026!", "Redirected to dashboard, sidebar visible"),
    ("Dashboard stats", "View dashboard after login", "Shows Projects, Credits, Evidence, Open Actions counts"),
    ("Navigation tabs", "Click each sidebar tab", "All 20 tabs load without errors"),
    ("Carbon Registry tab", "Click Carbon Registry", "Shows project list with status badges"),
    ("Credit Registry tab", "Click Credit Registry", "Shows credit batches with serial prefixes"),
    ("Anchoring tab", "Click Anchoring", "Shows Blockchain Anchoring panel with stats"),
    ("Run reconciliation", "Click Run Reconciliation", "Shows result (No anchors found, or list of anchors)"),
    ("Responsive layout", "Resize browser to 375px width", "Layout adapts, no horizontal scroll"),
    ("Logout", "Click Logout button", "Returns to login page, session cleared"),
]
for i, (action, cmd, expected) in enumerate(ui_tests, 1):
    add_test_row(ui_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.7 GIS Service Tests", level=2)
gis_table = doc.add_table(rows=1, cols=4)
gis_table.style = "Table Grid"
for i, h in enumerate(headers):
    gis_table.rows[0].cells[i].text = h
    gis_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(gis_table.rows[0].cells[i], "0f766e")
    gis_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

gis_tests = [
    ("GIS health", "curl http://localhost:8104/health", "Healthy response"),
    ("GIS assessment", "POST /api/v1/gis/assess/{project_id}", "Returns spatial analysis results"),
    ("Boundary validation", "POST boundary GeoJSON", "Returns validation status and area hectares"),
]
for i, (action, cmd, expected) in enumerate(gis_tests, 1):
    add_test_row(gis_table, i, f"{action}\n{cmd}", expected)

doc.add_paragraph()

doc.add_heading("9.8 Marketplace Tests", level=2)
market_table = doc.add_table(rows=1, cols=4)
market_table.style = "Table Grid"
for i, h in enumerate(headers):
    market_table.rows[0].cells[i].text = h
    market_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(market_table.rows[0].cells[i], "0f766e")
    market_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

market_tests = [
    ("List listings", "GET /api/v1/marketplace/listings", "Returns array of marketplace listings"),
    ("Create listing", "POST /api/v1/marketplace/listings", "Returns created listing with ID"),
    ("Execute trade", "POST /api/v1/marketplace/trades", "Returns trade confirmation"),
    ("Get prices", "GET /api/v1/marketplace/prices", "Returns current market prices"),
]
for i, (action, cmd, expected) in enumerate(market_tests, 1):
    add_test_row(market_table, i, f"{action}\n{cmd}", expected)

doc.add_page_break()

# =============================================================================
# 10. TROUBLESHOOTING
# =============================================================================
doc.add_heading("10. Troubleshooting", level=1)

doc.add_heading("10.1 Port Already in Use", level=2)
doc.add_paragraph("If you get EADDRINUSE errors, find and kill the process:")
add_code_block(doc, """# Windows:
netstat -ano | findstr :8082
taskkill //F //PID <PID>

# macOS/Linux:
lsof -i :8082
kill -9 <PID>""")

doc.add_heading("10.2 Database Connection Refused", level=2)
doc.add_paragraph("Ensure PostgreSQL is running and credentials are correct:")
add_code_block(doc, """# Check PostgreSQL status
pg_isready -h localhost -p 5432

# If using Docker:
docker ps | grep postgres""")

doc.add_heading("10.3 Python Module Not Found", level=2)
doc.add_paragraph("Ensure the virtual environment is activated:")
add_code_block(doc, """# Windows:
.venv\\Scripts\\activate

# Verify:
python -c 'import app.main'""")

doc.add_heading("10.4 Frontend Hydration Error", level=2)
doc.add_paragraph(
    "If the frontend shows hydration mismatch errors, this is a known Next.js 15 issue with dynamic imports. "
    "The current codebase includes a loading.tsx that returns null and a ClientOnly wrapper to prevent this. "
    "If issues persist, try clearing the Next.js cache:"
)
add_code_block(doc, "cd frontend/web-portal && rm -rf .next && npm run dev")

doc.add_heading("10.5 API Gateway 404 Errors", level=2)
doc.add_paragraph("If routes return 404 through the gateway but work directly:")
add_code_block(doc, """# Verify gateway is using updated routes:
curl http://localhost:8082/anchors/status/unanchored-count

# If 404, restart the gateway:
cd api-gateway && npm run dev""")

doc.add_page_break()

# =============================================================================
# 11. TEST SUMMARY CHECKLIST
# =============================================================================
doc.add_heading("11. Test Summary Checklist", level=1)
doc.add_paragraph("Use this checklist to verify all major features are working:")

checklist_table = doc.add_table(rows=1, cols=3)
checklist_table.style = "Table Grid"
headers = ["Category", "Test Count", "All Passed?"]
for i, h in enumerate(headers):
    checklist_table.rows[0].cells[i].text = h
    checklist_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(checklist_table.rows[0].cells[i], "0f766e")
    checklist_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

checklist = [
    ("Database Connectivity", "6", "[ ]"),
    ("Backend API Health", "6", "[ ]"),
    ("Authentication", "6", "[ ]"),
    ("Carbon Registry", "8", "[ ]"),
    ("Merkle Anchoring", "7", "[ ]"),
    ("Frontend UI", "10", "[ ]"),
    ("GIS Service", "3", "[ ]"),
    ("Marketplace", "4", "[ ]"),
    ("Total", "50", "[ ]"),
]
for cat, count, passed in checklist:
    row = checklist_table.add_row().cells
    row[0].text = cat
    row[1].text = count
    row[2].text = passed
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    if cat == "Total":
        for cell in row:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
doc.add_paragraph(
    "Document generated by DEPLOYMENT_AND_TESTING_GUIDE.docx.py  |  "
    "For issues, refer to README.md or contact the ZAI-CTS Platform Team."
)

# Save
doc.save("DEPLOYMENT_AND_TESTING_GUIDE.docx")
print("DOCX created: DEPLOYMENT_AND_TESTING_GUIDE.docx")
