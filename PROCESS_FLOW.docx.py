"""Generate ZAI-CTS Process Flow Document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_flow_box(doc, title, items, color="0f766e"):
    """Add a process box with title and bullet items."""
    p = doc.add_paragraph()
    run = p.add_run(f"▣ {title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    for item in items:
        doc.add_paragraph(f"  → {item}", style="List Bullet")
    doc.add_paragraph()

def add_decision(doc, question, yes_path, no_path):
    """Add a decision diamond description."""
    p = doc.add_paragraph()
    run = p.add_run(f"◆ DECISION: {question}")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
    doc.add_paragraph(f"  ✓ YES → {yes_path}")
    doc.add_paragraph(f"  ✗ NO  → {no_path}")
    doc.add_paragraph()

def add_arrow(doc, label=""):
    """Add a downward arrow."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"▼ {label}" if label else "▼")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    doc.add_paragraph()

doc = Document()

# Title
title = doc.add_heading("ZAI-CTS Process Flow Documentation", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0d, 0x94, 0x88)

subtitle = doc.add_paragraph("Complete system, business, and user process flows for the Zimbabwe AI-Enhanced Carbon Trading Ecosystem")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph("Generated: 2026-06-30")
doc.add_paragraph()

# =============================================================================
# 1. SYSTEM REQUEST FLOW
# =============================================================================
doc.add_heading("1. System Request Flow", level=1)
doc.add_paragraph(
    "This section describes how a single HTTP request travels through the ZAI-CTS architecture "
    "from the user's browser to the database and back."
)

add_flow_box(doc, "1. User Browser", [
    "User interacts with Next.js 15 web portal (http://localhost:3000)",
    "React components render MUI v6 interface with Tailwind CSS",
    "Client-side state managed via React hooks (useState, useEffect)",
    "Auth tokens stored in localStorage (access_token, refresh_token)",
])
add_arrow(doc, "HTTPS / HTTP")

add_flow_box(doc, "2. Next.js Frontend (Port 3000)", [
    "Page routing via Next.js App Router (app/ directory)",
    "ClientOnly wrapper prevents hydration mismatches",
    "Dynamic imports for heavy components (ssr: false)",
    "API calls made to API Gateway at localhost:8082",
    "Correlation ID (x-correlation-id) attached to every request",
])
add_arrow(doc, "HTTP Request")

add_flow_box(doc, "3. API Gateway (Fastify, Port 8082)", [
    "Receives request with auth headers (Bearer token, x-actor-id)",
    "CORS middleware allows localhost:3000 origin",
    "Rate limiting: 300 requests/minute per IP",
    "Helmet.js security headers applied",
    "Routes request to appropriate backend microservice:",
    "   /api/v1/auth/*, /api/v1/projects, /anchors/* → Carbon Registry (8102)",
    "   /api/v1/gis/* → GIS Service (8104)",
    "   /api/v1/marketplace/* → Marketplace (8106)",
    "   /api/v1/compliance/* → Compliance (8107)",
    "   /api/v1/ai/* → AI Validation (8103)",
])
add_arrow(doc, "HTTP Forward")

add_flow_box(doc, "4. Backend Microservice (FastAPI)", [
    "Receives request via Uvicorn ASGI server",
    "FastAPI dependency injection resolves AsyncSession from get_db_session",
    "JWT / session token validated (backend derives role from session)",
    "RBAC middleware checks permissions against required action",
    "Business logic executed in service/application layer",
    "Repository pattern abstracts database operations",
    "Audit events written to audit.audit_events table",
])
add_arrow(doc, "SQL Query")

add_flow_box(doc, "5. PostgreSQL Database (Port 5432)", [
    "SQLAlchemy async engine connects via asyncpg driver",
    "Query executed against appropriate schema (registry, identity, audit, etc.)",
    "PostGIS spatial queries use GEOMETRY/GIST indexes",
    "Transaction committed or rolled back",
    "Results serialized to Pydantic models",
])
add_arrow(doc, "Response Chain")

add_flow_box(doc, "6. Response Path", [
    "Database → Repository → Service → FastAPI → JSON response",
    "Gateway forwards response back to frontend",
    "Frontend renders updated UI state",
    "Loading indicators dismissed, success/error toasts shown",
])

doc.add_paragraph()
doc.add_paragraph("Complete Request Lifecycle:")
lifecycle = doc.add_paragraph()
lifecycle.add_run("Browser → Next.js → API Gateway (8082) → Microservice (8102-8107) → SQLAlchemy → PostgreSQL → Response → UI Update")
lifecycle.runs[0].font.size = Pt(9)
lifecycle.runs[0].font.name = "Consolas"
lifecycle.runs[0].font.color.rgb = RGBColor(0x0d, 0x94, 0x88)
doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# 2. CARBON CREDIT LIFECYCLE PROCESS FLOW
# =============================================================================
doc.add_heading("2. Carbon Credit Lifecycle Process Flow", level=1)
doc.add_paragraph(
    "The core business process from project conception to credit retirement. "
    "This is the primary value chain of the ZAI-CTS platform."
)

stages = [
    ("STAGE 1: Project Registration", "0f766e", [
        "Project Developer fills registration form (title, methodology, location, boundaries)",
        "System validates methodology against registry.methodologies table",
        "GIS boundary file uploaded (GeoJSON/KML/Shapefile)",
        "Project saved as 'draft' status in registry.carbon_projects",
        "Audit event: carbon.project.registered",
    ]),
    ("STAGE 2: Submission for Verification", "0f766e", [
        "Developer submits project for official verification",
        "Status changes to 'submitted_for_verification'",
        "Verification case automatically opened",
        "Required evidence checklist generated",
    ]),
    ("STAGE 3: Evidence Collection", "0f766e", [
        "Developer uploads evidence packages:",
        "   - Boundary files (GeoJSON/KML)",
        "   - Monitoring reports (PDF)",
        "   - Carbon calculations (Excel/CSV)",
        "   - Biomass inventory (CSV)",
        "   - Satellite imagery metadata (JSON/TIFF)",
        "   - Field photographs (JPG/PNG)",
        "   - Verifier statement (PDF)",
        "   - Digital signatures (PEM/SIG)",
        "Evidence stored in storage/evidence/ with checksums",
    ]),
    ("STAGE 4: Automatic Validation", "0f766e", [
        "System performs automated checks:",
        "   - File format validation",
        "   - Boundary geometry validation (PostGIS ST_IsValid)",
        "   - Serial number uniqueness check",
        "   - Calculation consistency checks",
        "Results stored in verification_cases table",
    ]),
    ("STAGE 5: AI Review", "0f766e", [
        "LangChain-powered AI validation service analyzes:",
        "   - PDD (Project Design Document) completeness",
        "   - Additionality assessment",
        "   - Leakage detection",
        "   - Legal compliance scan",
        "   - Price forecasting",
        "Confidence score (0-1) and risk level assigned",
    ]),
    ("STAGE 6: GIS Assessment", "0f766e", [
        "PostGIS spatial analysis executed:",
        "   - Boundary area calculation (hectares)",
        "   - Forest cover percentage",
        "   - Carbon density estimation (tCO2e/hectare)",
        "   - Fire risk assessment",
        "   - Centroid coordinates",
        "Satellite imagery compared with field evidence",
    ]),
    ("STAGE 7: Verifier Review", "0f766e", [
        "Accredited third-party verifier examines:",
        "   - Field inspection reports",
        "   - Monitoring data continuity",
        "   - Methodology compliance",
        "   - Carbon stock calculations",
        "Verifier signs with digital signature",
        "Decision: PASS / WARNING / FAIL / REQUEST_MORE_INFO",
    ]),
    ("STAGE 8: ZiCMA Regulatory Review", "0f766e", [
        "Zimbabwe Climate Authority (ZiCMA) reviews:",
        "   - National registry compliance",
        "   - NDC alignment",
        "   - Double-counting prevention",
        "   - Article 6 eligibility (if applicable)",
        "ZiCMA officer signs with digital signature",
        "Decision: APPROVE / REJECT / RETURN_FOR_REVISION",
    ]),
    ("STAGE 9: Credit Issuance", "0d9488", [
        "Upon approval, system issues serialized credits:",
        "   - Unique serial prefix generated (e.g., ZWE-RE-2025)",
        "   - Individual credit entries created with deterministic hashes",
        "   - Batch recorded in registry.carbon_credit_batches",
        "   - blockchain_tx_id stored (mock Fabric TX for now)",
        "Audit event: carbon.credits.issued",
    ]),
    ("STAGE 10: Anchoring (Blockchain Notarization)", "0d9488", [
        "Unanchored credit entries batched:",
        "   - SHA-256 hash computed for each entry",
        "   - Merkle tree built from entry hashes",
        "   - Merkle root stored in anchor_batches table",
        "   - Hash chain links to previous anchor",
        "   - Mock fabric_tx_id assigned (future: real Fabric TX)",
        "Provides tamper-evident audit trail",
    ]),
    ("STAGE 11: Marketplace Listing", "047857", [
        "Project owner creates marketplace listing:",
        "   - Quantity of credits for sale",
        "   - Price per tCO2e (USD)",
        "   - Vintage year",
        "   - Eligibility tags (CORSIA, VCS, etc.)",
        "Listing saved in market.marketplace_listings",
    ]),
    ("STAGE 12: Trade Execution", "047857", [
        "Buyer selects listing and places order:",
        "   - Quantity requested",
        "   - Payment method (bank_transfer, crypto, escrow)",
        "   - Buyer account ID",
        "Order created in market.marketplace_orders",
        "Status: pending_compliance_review",
    ]),
    ("STAGE 13: Compliance Review", "047857", [
        "Compliance service validates:",
        "   - Buyer KYC status",
        "   - Anti-money laundering (AML) check",
        "   - Export control verification",
        "   - Article 6 authorization (if international)",
        "Decision: APPROVED / REJECTED / MANUAL_REVIEW",
    ]),
    ("STAGE 14: Settlement", "047857", [
        "Upon compliance approval:",
        "   - Credits transferred from seller to buyer account",
        "   - Ledger updated in registry.registry_accounts",
        "   - Settlement record created",
        "   - Market settlement control recorded",
        "Audit event: carbon.credits.transferred",
    ]),
    ("STAGE 15: Retirement", "1e3a5f", [
        "Buyer retires credits for compliance:",
        "   - Retirement certificate generated",
        "   - Credits marked as 'retired' (non-transferable)",
        "   - Public disclosure published",
        "   - NDC accounting snapshot updated",
        "   - Corresponding adjustment applied (Article 6)",
        "Audit event: carbon.credits.retired",
    ]),
]

for stage_title, color, items in stages:
    add_flow_box(doc, stage_title, items, color)
    if stage_title != "STAGE 15: Retirement":
        add_arrow(doc)

doc.add_page_break()

# =============================================================================
# 3. USER ROLE PROCESS FLOWS
# =============================================================================
doc.add_heading("3. User Role Process Flows", level=1)
doc.add_paragraph("Each user role has distinct permissions and workflows within the platform.")

roles = [
    ("3.1 Project Developer Flow", [
        "1. Register new organization (if not exists)",
        "2. Create user account with 'project_developer' role",
        "3. Wait for ZiCMA approval of account",
        "4. Login to portal",
        "5. Navigate to Carbon Registry → Register Project",
        "6. Fill project details (title, methodology, district, province)",
        "7. Upload boundary evidence (GeoJSON/KML)",
        "8. Submit project for verification",
        "9. Upload verification evidence package",
        "10. Respond to verifier queries",
        "11. Receive approval notification",
        "12. View issued credits in Credit Registry",
        "13. Create marketplace listing",
        "14. Monitor sales and transfers",
    ]),
    ("3.2 ZiCMA Administrator Flow", [
        "1. Login as seed administrator (pre-created)",
        "2. Review pending user registrations",
        "3. Approve/reject/suspend user accounts",
        "4. Review submitted projects",
        "5. Advance project workflow stages",
        "6. Issue credit batches upon verification completion",
        "7. Monitor national readiness dashboard",
        "8. Execute national stage controls (rulebook adoption, methodology approval)",
        "9. Open compliance cases",
        "10. Review audit timeline",
        "11. Generate national accounting snapshots",
        "12. Oversee anchoring operations",
    ]),
    ("3.3 Verifier Flow", [
        "1. Register with 'verifier' role",
        "2. Wait for ZiCMA accreditation grant",
        "3. Login to portal",
        "4. View assigned verification cases",
        "5. Review evidence packages",
        "6. Conduct field inspections (offline)",
        "7. Upload verifier statement (PDF)",
        "8. Submit verification decision (PASS/WARNING/FAIL)",
        "9. Sign with digital signature",
        "10. View verification history",
    ]),
    ("3.4 Corporate Buyer Flow", [
        "1. Register with 'corporate_buyer' role",
        "2. Open registry account",
        "3. Complete KYC verification",
        "4. Login to portal",
        "5. Browse marketplace listings",
        "6. Filter by vintage, price, methodology",
        "7. Place order on selected listing",
        "8. Complete payment (bank transfer/crypto)",
        "9. Receive credit transfer to account",
        "10. View portfolio in Credit Registry",
        "11. Retire credits for compliance",
        "12. Download retirement certificate",
    ]),
    ("3.5 GIS Analyst Flow", [
        "1. Login with 'gis_analyst' role",
        "2. Navigate to GIS Intelligence tab",
        "3. Run spatial assessment on project boundaries",
        "4. Upload satellite imagery metadata",
        "5. Validate boundary against national GIS layers",
        "6. Submit GIS validation decision",
        "7. Record GIS processing jobs",
    ]),
    ("3.6 Compliance Officer Flow", [
        "1. Login with 'compliance_officer' role",
        "2. Monitor marketplace for suspicious activity",
        "3. Review trade settlements",
        "4. Open enforcement cases",
        "5. Conduct fraud risk assessment",
        "6. Freeze suspicious credits",
        "7. Generate compliance reports",
    ]),
]

for role_title, steps in roles:
    doc.add_heading(role_title, level=2)
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph()

doc.add_page_break()

# =============================================================================
# 4. MERKLE ANCHORING PROCESS FLOW
# =============================================================================
doc.add_heading("4. Merkle Root Anchoring Process Flow", level=1)
doc.add_paragraph(
    "The blockchain anchoring pattern (Option B) provides tamper-evident record keeping "
    "without requiring a live Hyperledger Fabric network."
)

doc.add_heading("4.1 Credit Entry Creation Flow", level=2)
anchor_flow_1 = [
    ("1. Credit Batch Issued", [
        "ZiCMA admin issues credits for approved project",
        "System generates unique serial numbers (e.g., ZWE-RE-2025-00001)",
        "Credits saved in registry.carbon_credit_batches",
    ]),
    ("2. Individual Entries Created", [
        "Each credit becomes a row in registry.carbon_credit_entries",
        "Deterministic SHA-256 hash computed from:",
        "   - serial_number + vintage_year + quantity_tco2e + project_id + batch_id",
        "entry_hash stored (64-char hex)",
        "anchor_batch_id is NULL (unanchored)",
    ]),
    ("3. Unanchored Count Tracked", [
        "Frontend displays count of unanchored entries",
        "System monitors via GET /anchors/status/unanchored-count",
    ]),
]
for title, items in anchor_flow_1:
    add_flow_box(doc, title, items)
    add_arrow(doc)

doc.add_heading("4.2 Anchor Batch Creation Flow", level=2)
anchor_flow_2 = [
    ("4. Trigger Anchor Batch", [
        "Admin clicks 'Anchor Now' in Anchoring tab",
        "Frontend sends POST /anchors/batch",
        "API Gateway forwards to Carbon Registry (8102)",
    ]),
    ("5. Fetch Unanchored Entries", [
        "AnchoringService queries carbon_credit_entries WHERE anchor_batch_id IS NULL",
        "If count = 0 → return 400 'no_entries'",
    ]),
    ("6. Compute Entry Hashes", [
        "For each entry, recompute deterministic hash",
        "Hashes collected into array [h1, h2, h3, ..., hn]",
    ]),
    ("7. Build Merkle Tree", [
        "compute_merkle_root() builds binary hash tree:",
        "   Level 0: [h1, h2, h3, h4, h5, h6, h7, h8]",
        "   Level 1: [H(h1+h2), H(h3+h4), H(h5+h6), H(h7+h8)]",
        "   Level 2: [H(H12+H34), H(H56+H78)]",
        "   Level 3 (Root): H(H1234+H5678)",
        "If odd number of leaves, last hash is duplicated",
    ]),
    ("8. Build Hash Chain", [
        "Fetch previous anchor (latest anchored batch)",
        "Compute anchor_hash = SHA256(merkle_root + previous_anchor_hash + timestamp + entry_count)",
        "This creates an immutable linked list of anchors",
    ]),
    ("9. Store Anchor Record", [
        "Create row in registry.anchor_batches:",
        "   - batch_name (auto-generated or custom)",
        "   - from_record_id, to_record_id (entry range)",
        "   - entry_count",
        "   - merkle_root (64-char hex)",
        "   - previous_anchor_id, previous_anchor_hash",
        "   - fabric_tx_id (mock: 'fabric:<uuid>')",
        "   - fabric_block_number (auto-increment)",
        "   - anchored_at (timestamp)",
        "   - status = 'anchored'",
    ]),
    ("10. Link Entries to Anchor", [
        "Update all anchored entries: SET anchor_batch_id = new_anchor_id",
        "Commit transaction (all-or-nothing)",
    ]),
    ("11. Return Result", [
        "Response includes:",
        "   - anchor_id, batch_name, merkle_root",
        "   - anchor_hash, entry_count",
        "   - fabric_tx_id, fabric_block_number",
        "   - previous_anchor_id",
    ]),
]
for title, items in anchor_flow_2:
    add_flow_box(doc, title, items)
    add_arrow(doc)

doc.add_heading("4.3 Verification & Reconciliation Flow", level=2)
anchor_flow_3 = [
    ("12. Verify Single Anchor", [
        "User enters anchor_id and clicks 'Verify'",
        "GET /anchors/{anchor_id}/verify",
        "System:",
        "   a. Fetches anchor record",
        "   b. Fetches all linked credit entries",
        "   c. Recomputes entry hashes",
        "   d. Rebuilds Merkle tree",
        "   e. Compares recomputed_root == stored_root",
        "Returns: 'verified' or 'tampered'",
    ]),
    ("13. Reconcile Entire Chain", [
        "User clicks 'Run Reconciliation'",
        "GET /anchors/reconcile",
        "System:",
        "   a. Fetches ALL anchors ordered by anchored_at",
        "   b. Verifies each anchor individually (step 12)",
        "   c. Checks chain continuity:",
        "      anchor[n].previous_anchor_hash == anchor[n-1].merkle_root",
        "   d. Flags any breaks in the chain",
        "Returns: Array of verification results with chain_continuous flag",
    ]),
    ("14. Tamper Detection", [
        "If any entry is modified in database:",
        "   - Recomputed hash ≠ stored hash",
        "   - Merkle root mismatch detected",
        "   - Verification returns 'tampered'",
        "   - Chain reconciliation flags the break",
        "Provides cryptographic proof of data integrity",
    ]),
]
for title, items in anchor_flow_3:
    add_flow_box(doc, title, items)

doc.add_page_break()

# =============================================================================
# 5. AUTHENTICATION & SESSION FLOW
# =============================================================================
doc.add_heading("5. Authentication & Session Flow", level=1)

auth_flow = [
    ("1. User Registration", [
        "User fills registration form (full_name, email, password, role)",
        "System generates salt (32 random bytes)",
        "Password hashed with PBKDF2-HMAC-SHA256 (100,000 iterations)",
        "User saved to identity.users with status='pending_approval'",
        "Email verification token generated",
    ]),
    ("2. Account Approval", [
        "ZiCMA admin reviews pending users",
        "Approves account → status='approved'",
        "Rejects account → status='rejected' + reason",
        "Suspends account → status='suspended'",
    ]),
    ("3. User Login", [
        "User submits email + password",
        "System fetches user by email",
        "Recomputes password hash with stored salt",
        "Compares hash == stored_hash",
        "If MFA enabled → prompt for TOTP code",
    ]),
    ("4. Session Creation", [
        "On successful authentication:",
        "   - JWT access_token generated (15-min expiry)",
        "   - Refresh token generated (7-day expiry)",
        "   - Session record saved to identity.sessions",
        "   - token_hash (SHA-256 of token) stored",
        "   - IP address and device fingerprint recorded",
    ]),
    ("5. Authenticated Request", [
        "Frontend sends Bearer token in Authorization header",
        "Backend validates JWT signature and expiry",
        "Derives user role from session (NOT from frontend header)",
        "Security fix: x-actor-role is ignored; backend looks up role",
        "RBAC check: does user have permission for this action?",
    ]),
    ("6. Token Refresh", [
        "Access token expires after 15 minutes",
        "Frontend sends refresh_token to /auth/refresh",
        "System validates refresh token",
        "Issues new access_token + refresh_token pair",
        "Old refresh token invalidated (rotation)",
    ]),
    ("7. Logout", [
        "User clicks Logout",
        "Frontend clears localStorage tokens",
        "Backend invalidates session in database",
        "Refresh token added to revocation list",
    ]),
]

for title, items in auth_flow:
    add_flow_box(doc, title, items)
    add_arrow(doc)

doc.add_page_break()

# =============================================================================
# 6. AUDIT & COMPLIANCE FLOW
# =============================================================================
doc.add_heading("6. Audit & Compliance Flow", level=1)

doc.add_heading("6.1 Automatic Audit Logging", level=2)
audit_auto = [
    "Every authenticated action triggers an audit event:",
    "   - event_type (e.g., carbon.project.registered)",
    "   - actor_id (UUID of user)",
    "   - actor_role (derived from session)",
    "   - organization_id",
    "   - resource_type + resource_id",
    "   - action (create, read, update, delete)",
    "   - outcome (success, failure)",
    "   - ip_address + device fingerprint",
    "   - old_value + new_value (JSONB diffs)",
    "   - correlation_id (links related actions)",
    "   - metadata (additional context)",
    "Events saved to audit.audit_events table",
    "Indexed by (resource_type, resource_id) and correlation_id",
]
for item in audit_auto:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph()
doc.add_heading("6.2 Compliance Investigation Flow", level=2)
compliance_flow = [
    ("1. Trigger Detection", [
        "Automated alerts detect anomalies:",
        "   - Unusual trading volume",
        "   - Rapid credit transfers",
        "   - Failed verification patterns",
        "   - Geographic anomalies",
    ]),
    ("2. Case Opening", [
        "Compliance officer opens case",
        "Case saved to compliance.compliance_cases",
        "Linked to affected credits and accounts",
    ]),
    ("3. Evidence Collection", [
        "Officer reviews:",
        "   - Full audit trail (all related events)",
        "   - Transaction history",
        "   - User activity logs",
        "   - External data sources",
    ]),
    ("4. Decision", [
        "Officer makes determination:",
        "   - No violation → Case closed",
        "   - Minor violation → Warning issued",
        "   - Major violation → Credits frozen, account suspended",
        "   - Criminal activity → Referred to authorities",
    ]),
    ("5. Remediation", [
        "If credits frozen:",
        "   - freezeLedgerCredits() called",
        "   - Frozen credits cannot be transferred/retired",
        "   - Public disclosure may be required",
    ]),
]
for title, items in compliance_flow:
    add_flow_box(doc, title, items)
    add_arrow(doc)

doc.add_page_break()

# =============================================================================
# 7. DATA FLOW DIAGRAM (TABLE FORMAT)
# =============================================================================
doc.add_heading("7. Data Flow Summary", level=1)

data_flow_table = doc.add_table(rows=1, cols=5)
data_flow_table.style = "Table Grid"
headers = ["Source", "Action", "Data", "Destination", "Result"]
for i, h in enumerate(headers):
    data_flow_table.rows[0].cells[i].text = h
    data_flow_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(data_flow_table.rows[0].cells[i], "0f766e")
    data_flow_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

flows = [
    ("Browser", "POST /auth/login", "{email, password}", "API Gateway", "{access_token, user}"),
    ("API Gateway", "Forward", "Headers + Body", "Carbon Registry", "Process auth"),
    ("Carbon Registry", "SELECT", "WHERE email=", "PostgreSQL", "User record"),
    ("Carbon Registry", "INSERT", "Session tokens", "PostgreSQL", "Session saved"),
    ("Carbon Registry", "Response", "JWT + user", "API Gateway", "Forward response"),
    ("API Gateway", "Response", "JSON", "Browser", "Store tokens, redirect"),
    ("Browser", "POST /projects", "Project data", "API Gateway", "Forward to registry"),
    ("Carbon Registry", "INSERT", "Project record", "PostgreSQL", "Project created"),
    ("Carbon Registry", "INSERT", "Audit event", "PostgreSQL", "Audit logged"),
    ("Carbon Registry", "Response", "Project JSON", "API Gateway", "Forward"),
    ("Browser", "POST /anchors/batch", "Batch name", "API Gateway", "Forward"),
    ("Carbon Registry", "SELECT", "Unanchored entries", "PostgreSQL", "Entry hashes"),
    ("Carbon Registry", "Compute", "Merkle tree", "Memory", "Merkle root"),
    ("Carbon Registry", "INSERT", "Anchor batch", "PostgreSQL", "Anchor saved"),
    ("Carbon Registry", "UPDATE", "Entry anchor IDs", "PostgreSQL", "Entries linked"),
    ("Carbon Registry", "Response", "Anchor result", "API Gateway", "Forward"),
]

for src, action, data, dest, result in flows:
    row = data_flow_table.add_row().cells
    row[0].text = src
    row[1].text = action
    row[2].text = data
    row[3].text = dest
    row[4].text = result
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph(
    "Document generated by PROCESS_FLOW.docx.py  |  "
    "For detailed API specifications, see docs/api/ and docs/ai/"
)

# Save
doc.save("PROCESS_FLOW.docx")
print("DOCX created: PROCESS_FLOW.docx")
