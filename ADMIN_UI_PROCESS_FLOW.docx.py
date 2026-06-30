"""Generate ZAI-CTS Admin-Only UI Process Flow Document for Presentation."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_ui_step(doc, step_num, title, screen, actions, ui_shows, outcome, tip=""):
    """Add a structured UI process step."""
    p = doc.add_paragraph()
    run = p.add_run(f"STEP {step_num}: {title}")
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0d, 0x94, 0x88)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)
    
    cells = table.rows[0].cells
    cells[0].text = "🖥️ Screen"
    cells[1].text = screen
    set_cell_shading(cells[0], "f1f5f9")
    
    cells = table.rows[1].cells
    cells[0].text = "👆 Admin Actions"
    cells[1].text = actions
    set_cell_shading(cells[0], "f1f5f9")
    
    cells = table.rows[2].cells
    cells[0].text = "👁️ UI Shows"
    cells[1].text = ui_shows
    set_cell_shading(cells[0], "f1f5f9")
    
    cells = table.rows[3].cells
    cells[0].text = "✅ Outcome"
    cells[1].text = outcome
    set_cell_shading(cells[0], "dcfce7")
    
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if cell.text in ["🖥️ Screen", "👆 Admin Actions", "👁️ UI Shows", "✅ Outcome"]:
                cell.paragraphs[0].runs[0].font.bold = True
    
    if tip:
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Tip: {tip}")
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
    
    doc.add_paragraph()

def add_phase_header(doc, phase_num, phase_title, color="0f766e"):
    """Add a colored phase header."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"▓▓▓  PHASE {phase_num}: {phase_title}  ▓▓▓")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(p._p.get_or_add_pPr().get_or_add_shd(), color) if hasattr(p._p, 'get_or_add_pPr') else None
    # Alternative: just color the text
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    doc.add_paragraph()


doc = Document()

# Title
title = doc.add_heading("ZAI-CTS Admin UI Process Flow", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0d, 0x94, 0x88)

subtitle = doc.add_paragraph("Complete step-by-step guide for a ZiCMA Administrator to process a carbon project from registration to retirement")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph("Role: ZiCMA Registry Administrator (Super Administrator)  |  Login: admin@zai-cts.gov.zw / ZAI-CTS-Admin-2026!", style="Intense Quote")
doc.add_paragraph("Generated: 2026-06-30")
doc.add_paragraph()

# =============================================================================
# PHASE 1: LOGIN & SETUP
# =============================================================================
doc.add_heading("PHASE 1: Login & Platform Access", level=1)
set_cell_shading(doc.add_paragraph()._p.get_or_add_pPr(), "0f766e") if False else None

add_ui_step(doc, 1, "Navigate to Portal & Login", 
    "Login Page (http://localhost:3000)",
    "• Open browser and navigate to http://localhost:3000\n• Enter Email: admin@zai-cts.gov.zw\n• Enter Password: ZAI-CTS-Admin-2026!\n• Click 'Login' button",
    "• ZAI-CTS logo and loading animation\n• Login form with pre-filled credentials\n• Alert banner: 'Seed administrator: admin@zai-cts.gov.zw'\n• MFA code field (optional)",
    "✓ Redirected to Operations Portal dashboard\n✓ Sidebar shows 21 navigation tabs\n✓ Top bar shows: 'ZiCMA Registry Administrator - Super Administrator'\n✓ Gateway health indicator: 'healthy'",
    "If port 3000 is busy, the app auto-switches to 3004")

add_ui_step(doc, 2, "Verify System Health", 
    "Dashboard Tab (default view)",
    "• Observe the dashboard stats cards\n• Click 'Refresh' button in top bar\n• Review 'Verification operations overview' section",
    "• 4 stat cards: Projects (N), Credits (N), Evidence (N), Open Actions (N)\n• 'Outstanding actions' panel\n• 'Latest audit' panel with recent events\n• Green health indicator",
    "✓ All services responsive\n✓ Gateway healthy\n✓ Dashboard data loaded from PostgreSQL",
    "Check that Gateway shows 'healthy' before proceeding")

doc.add_page_break()

# =============================================================================
# PHASE 2: PROJECT REGISTRATION
# =============================================================================
doc.add_heading("PHASE 2: Project Registration", level=1)

add_ui_step(doc, 3, "Navigate to Carbon Registry", 
    "Carbon Registry Tab",
    "• Click 'Carbon Registry' in left sidebar\n• Review existing projects list\n• Click 'Register New Project' button",
    "• Table of existing carbon projects\n• Columns: Code, Title, Methodology, District, Province, Status\n• Status badges: draft, submitted_for_verification, under_verification, approved, credits_issued\n• 'Register New Project' button (top right)",
    "✓ Carbon Registry workspace loaded\n✓ Project list fetched from /api/v1/projects",
    "Use the search box to filter existing projects")

add_ui_step(doc, 4, "Fill Project Registration Form", 
    "Project Registration Form",
    "• Project Code: Enter unique code (e.g., RE-20250010)\n• Title: Enter project name (e.g., 'Kariba REDD+ Expansion 2025')\n• Description: Enter detailed project narrative\n• Methodology: Select from dropdown (e.g., VM0015)\n• District: Enter district name\n• Province: Enter province name\n• Estimated Annual tCO2e: Enter numeric value\n• Start Date: Select from date picker\n• Crediting Period (Years): Enter integer\n• Click 'Submit' button",
    "• Multi-field form with validation\n• Dropdown for methodology selection\n• Numeric validation on tCO2e field\n• Date picker for start_date\n• Required field indicators (*)",
    "✓ Project saved to registry.carbon_projects\n✓ Status set to 'draft'\n✓ Audit event logged: carbon.project.registered\n✓ Success toast notification",
    "Project code must be unique; system will reject duplicates")

add_ui_step(doc, 5, "View New Project in List", 
    "Carbon Registry Tab (updated)",
    "• Scroll through project list\n• Click on the newly created project row",
    "• New project appears in table with 'draft' status badge (gray)\n• Project code, title, and district visible\n• Clickable row for details",
    "✓ Project visible in registry\n✓ Status confirmed as 'draft'",
    "The project is now selectable for further actions")

doc.add_page_break()

# =============================================================================
# PHASE 3: PROJECT LIFECYCLE ADVANCEMENT
# =============================================================================
doc.add_heading("PHASE 3: Project Lifecycle Advancement", level=1)

add_ui_step(doc, 6, "Submit Project for Verification", 
    "Project Lifecycle Tab",
    "• Click 'Project Lifecycle' in sidebar\n• Select the project from dropdown (e.g., RE-20250010)\n• Click 'Submit for Verification' button",
    "• Workflow stepper: Draft → Submitted → Verification → Approved → Credits Issued\n• Current step highlighted: 'Draft'\n• Project selector dropdown\n• Action buttons for each workflow stage",
    "✓ Status changes to 'submitted_for_verification'\n✓ Workflow stepper advances to step 2\n✓ Verification case automatically opened\n✓ Audit event: carbon.project.submit_for_verification",
    "This triggers the entire verification pipeline")

doc.add_page_break()

# =============================================================================
# PHASE 4: EVIDENCE UPLOAD
# =============================================================================
doc.add_heading("PHASE 4: Evidence Upload & Validation", level=1)

add_ui_step(doc, 7, "Upload Verification Evidence", 
    "Verification Tab",
    "• Click 'Verification' in sidebar\n• Select project from dropdown\n• Review REQUIRED_VERIFICATION_UPLOADS checklist\n• Drag & drop or click to upload files:\n  - Boundary file (GeoJSON/KML/SHP)\n  - Monitoring report (PDF)\n  - Carbon calculations (XLSX/CSV)\n  - Biomass inventory (CSV)\n  - Satellite imagery metadata (JSON/TIFF)\n  - Field photographs (JPG/PNG)\n  - Verifier statement (PDF)\n  - Digital signature (PEM/SIG)\n• Click 'Upload Evidence Package'",
    "• Verification sequence stepper:\n  pending_evidence → evidence_uploaded → automatic_validation → ai_review → gis_review → mrv_review → verifier_review → zicma_review → approved → credit_issued\n• File drop zone with category auto-detection\n• Format validation indicators (green check / red X)\n• Progress bar during upload",
    "✓ Evidence saved to storage/evidence/\n✓ Verification case updated: status = 'evidence_uploaded'\n✓ Each file linked to category and digital signature\n✓ Checksum computed for integrity",
    "Files are automatically categorized by extension and filename keywords")

add_ui_step(doc, 8, "Run Automatic Validation", 
    "Verification Tab",
    "• Click 'Run Automatic Validation' button\n• Wait for processing",
    "• Loading spinner on button\n• Progress messages:\n  'Validating file formats...'\n  'Checking boundary geometry...'\n  'Verifying serial uniqueness...'\n  'Cross-referencing calculations...'",
    "✓ Validation results displayed:\n  - Format status: valid / requires_review\n  - Boundary validation: valid / invalid\n  - Serial check: unique / duplicate_found\n  - Calculation consistency: passed / failed\n✓ Verification case: status = 'automatic_validation'",
    "Any 'requires_review' files must be addressed before proceeding")

doc.add_page_break()

# =============================================================================
# PHASE 5: AI & GIS REVIEWS
# =============================================================================
doc.add_heading("PHASE 5: AI Review & GIS Assessment", level=1)

add_ui_step(doc, 9, "Run AI Validation Review", 
    "AI Intelligence Tab",
    "• Click 'AI Intelligence' in sidebar\n• Select project from dropdown\n• Click 'Run AI Review' button",
    "• AI processing indicator with animated dots\n• Model info: model_version, prompt_version\n• Progress stages:\n  'Analyzing PDD completeness...'\n  'Assessing additionality...'\n  'Detecting leakage risks...'\n  'Scanning legal compliance...'",
    "✓ AI Review report generated:\n  - confidence_score: 0.00 - 1.00\n  - risk_level: low / medium / high / critical\n  - findings: array of structured observations\n  - recommendation: approve / review / reject\n✓ Verification case: status = 'ai_review'",
    "High confidence (>0.8) and low risk = fast-track approval")

add_ui_step(doc, 10, "Run GIS Assessment", 
    "GIS Intelligence Tab",
    "• Click 'GIS Intelligence' in sidebar\n• Select project from dropdown\n• Click 'Run GIS Assessment' button",
    "• Map visualization area (if enabled)\n• Spatial analysis progress:\n  'Computing boundary area...'\n  'Analyzing forest cover...'\n  'Estimating carbon density...'\n  'Assessing fire risk...'",
    "✓ GIS Assessment report:\n  - estimated_area_hectares\n  - forest_cover_percent\n  - carbon_density_tco2e_per_hectare\n  - fire_risk_level: low / medium / high\n  - boundary_validation_status: valid / invalid\n  - centroid_latitude / centroid_longitude\n✓ Verification case: status = 'gis_review'",
    "Invalid boundary will block project approval")

doc.add_page_break()

# =============================================================================
# PHASE 6: REGULATORY DECISIONS
# =============================================================================
doc.add_heading("PHASE 6: Regulatory Review & Approval", level=1)

add_ui_step(doc, 11, "Make Verifier Decision", 
    "Verification Tab → Verifier Stage",
    "• Return to 'Verification' tab\n• Review all evidence, AI findings, and GIS report\n• Select decision from dropdown:\n  - PASS\n  - WARNING\n  - FAIL\n  - REQUEST_MORE_INFORMATION\n• Enter comments in text area\n• Click 'Submit Verifier Decision'",
    "• Collapsible panels for each evidence category\n• AI findings summary card\n• GIS assessment summary card\n• Decision dropdown with color coding\n• Digital signature auto-generated: SIG-VERIFIER-<UUID>",
    "✓ Verification case: status = 'verifier_review'\n✓ Decision recorded with timestamp\n✓ Digital signature stored\n✓ Audit event: verification.verifier.decision",
    "WARNING decision requires follow-up; FAIL blocks the project")

add_ui_step(doc, 12, "Make ZiCMA Final Approval", 
    "Verification Tab → ZiCMA Stage",
    "• Review verifier decision and all supporting evidence\n• Select decision from dropdown:\n  - APPROVE\n  - REJECT\n  - RETURN_FOR_REVISION\n• Enter regulatory comments\n• Click 'Submit ZiCMA Decision'",
    "• Complete audit trail visible\n• All previous decisions summarized\n• NDC alignment indicator\n• Article 6 eligibility flag\n• Double-counting prevention check",
    "✓ If APPROVED:\n  - Project status = 'approved'\n  - Workflow stepper advances to 'Approved'\n  - Ready for credit issuance\n✓ If REJECTED:\n  - Project status = 'rejected'\n  - Appeal process initiated\n✓ Audit event: verification.zicma.decision",
    "RETURN_FOR_REVISION sends project back to developer for corrections")

doc.add_page_break()

# =============================================================================
# PHASE 7: CREDIT ISSUANCE
# =============================================================================
doc.add_heading("PHASE 7: Credit Issuance", level=1)

add_ui_step(doc, 13, "Issue Credit Batch", 
    "Credit Registry Tab",
    "• Click 'Credit Registry' in sidebar\n• Select the approved project\n• Click 'Issue Credits' button\n• Enter vintage year (e.g., 2025)\n• Enter quantity in tCO2e (e.g., 15000.0000)\n• System auto-generates serial prefix (e.g., ZWE-KA-2025)\n• Click 'Confirm Issuance'",
    "• Credit batches table for selected project\n• Issuance form with numeric inputs\n• Serial prefix preview\n• Confirmation dialog with summary",
    "✓ Credit batch created in registry.carbon_credit_batches\n✓ Status = 'issued'\n✓ blockchain_tx_id stored (mock Fabric TX)\n✓ Individual credit entries created in carbon_credit_entries\n✓ Each entry has deterministic SHA-256 hash\n✓ Audit event: carbon.credits.issued",
    "Serial prefix format: ZWE-{district_code}-{vintage_year}-{sequence}")

add_ui_step(doc, 14, "View Issued Credits", 
    "Credit Registry Tab (updated)",
    "• Review the newly issued batch\n• Click on batch to expand details\n• Review individual serial numbers",
    "• Batch card with:\n  - Vintage year\n  - Quantity tCO2e\n  - Serial prefix\n  - Status badge (green = issued)\n  - Blockchain TX ID (shortened with tooltip)\n• Expandable list of individual credits",
    "✓ Credits visible in registry\n✓ Each credit has unique serial number\n✓ All entries show 'unanchored' status",
    "Credits must be anchored before marketplace listing")

doc.add_page_break()

# =============================================================================
# PHASE 8: BLOCKCHAIN ANCHORING
# =============================================================================
doc.add_heading("PHASE 8: Blockchain Anchoring", level=1)

add_ui_step(doc, 15, "Navigate to Anchoring Tab", 
    "Anchoring Tab",
    "• Click 'Anchoring' in left sidebar (bottom of list)\n• Review unanchored credit count",
    "• 'Blockchain Anchoring' header\n• Shield icon with unanchored count (large number)\n• Description: 'Merkle root anchoring keeps a tamper-evident chain of credit hashes inside PostgreSQL'\n• 3 action panels:",
    "✓ Anchoring dashboard loaded\n✓ Unanchored count fetched from /anchors/status/unanchored-count",
    "The Anchoring tab is only visible to admin and compliance roles")

add_ui_step(doc, 16, "Create Anchor Batch", 
    "Anchoring Tab → Create Anchor Batch",
    "• Enter optional batch name (e.g., 'Kariba-2025-Q2')\n• Click 'Anchor Now' button",
    "• Text field for batch name\n• 'Anchor Now' button (disabled if count = 0)\n• Loading spinner during processing",
    "✓ Anchor batch created:\n  - anchor_id: UUID\n  - batch_name: user-provided or auto-generated\n  - merkle_root: 64-char SHA-256 hex\n  - anchor_hash: chain hash including previous anchor\n  - entry_count: number of credits anchored\n  - fabric_tx_id: mock Fabric transaction ID\n  - fabric_block_number: auto-incremented integer\n  - previous_anchor_id: links to previous batch (or null for genesis)",
    "Each anchor batch is immutable once created")

add_ui_step(doc, 17, "Verify Anchor Integrity", 
    "Anchoring Tab → Verify Anchor",
    "• Copy the anchor_id from the result card\n• Paste into 'Anchor ID' field\n• Click 'Verify' button",
    "• Input field for anchor ID\n• 'Verify' button\n• Result alert:\n  Green (success): 'Anchor verified' with stored vs recomputed root\n  Red (error): 'Anchor tampered' with root mismatch details",
    "✓ System recomputes Merkle root from database entries\n✓ Compares with stored merkle_root\n✓ Returns is_valid = true / false",
    "Verification proves data integrity without modifying anything")

add_ui_step(doc, 18, "Reconcile Anchor Chain", 
    "Anchoring Tab → Reconcile Chain",
    "• Click 'Run Reconciliation' button",
    "• 'Run Reconciliation' button\n• LinearProgress bar during processing\n• Results list of all anchors:\n  - Anchor 1: Valid ✓ Chain continuous ✓\n  - Anchor 2: Valid ✓ Chain continuous ✓\n  - Anchor N: ...\n• Each card shows:\n  - Valid / Tampered badge\n  - Chain Break warning (if any)\n  - Entry count\n  - Fabric TX ID\n  - Shortened Merkle root (hover for full)",
    "✓ Every anchor verified individually\n✓ Hash chain continuity checked:\n  anchor[n].previous_anchor_hash == anchor[n-1].merkle_root\n✓ Any breaks flagged with warning badge",
    "Run reconciliation weekly as part of audit procedures")

doc.add_page_break()

# =============================================================================
# PHASE 9: MARKETPLACE OPERATIONS
# =============================================================================
doc.add_heading("PHASE 9: Marketplace Listing & Trading", level=1)

add_ui_step(doc, 19, "Create Marketplace Listing", 
    "Marketplace Tab",
    "• Click 'Marketplace' in sidebar\n• Click 'Create Listing' button\n• Select project from dropdown\n• Enter quantity (tCO2e)\n• Enter price per tCO2e (USD)\n• Select vintage year\n• Add eligibility tags (CORSIA, VCS, Gold Standard, etc.)\n• Click 'Publish Listing'",
    "• Marketplace listings table\n• Create listing form\n• Price input with currency selector\n• Multi-select eligibility tags\n• Preview card before publishing",
    "✓ Listing saved to market.marketplace_listings\n✓ Status = 'open'\n✓ Visible to all buyers\n✓ Audit event: marketplace.listing.created",
    "Price is locked at listing creation; edits require new listing")

add_ui_step(doc, 20, "Execute Trade", 
    "Marketplace Tab → Trade Execution",
    "• Review open listing\n• Click 'Execute Trade' on a listing\n• Select buyer account from dropdown\n• Enter quantity to purchase\n• Select payment method (bank_transfer / crypto / escrow)\n• Click 'Confirm Trade'",
    "• Listing detail card\n• Buyer account selector\n• Quantity input (max = listing quantity)\n• Payment method radio buttons\n• Trade summary with total cost",
    "✓ Order created in market.marketplace_orders\n✓ Status = 'pending_compliance_review'\n✓ Credits reserved (not yet transferred)",
    "Trade requires compliance approval before settlement")

doc.add_page_break()

# =============================================================================
# PHASE 10: COMPLIANCE & SETTLEMENT
# =============================================================================
doc.add_heading("PHASE 10: Compliance Review & Settlement", level=1)

add_ui_step(doc, 21, "Review Trade for Compliance", 
    "Compliance Tab",
    "• Click 'Compliance' in sidebar\n• Review pending trades list\n• Click on trade to expand details\n• Review buyer KYC status\n• Check AML indicators\n• Verify export controls (if international)\n• Select decision:\n  - APPROVED\n  - REJECTED\n  - MANUAL_REVIEW\n• Click 'Submit Compliance Decision'",
    "• Pending trades table\n• Buyer profile summary\n• Risk indicators (low / medium / high)\n• KYC status badge\n• Decision dropdown with comments field",
    "✓ If APPROVED:\n  - Order status = 'approved'\n  - Proceeds to settlement\n✓ If REJECTED:\n  - Order cancelled\n  - Credits released back to seller\n✓ Audit event: compliance.trade.reviewed",
    "International trades require Article 6 authorization check")

add_ui_step(doc, 22, "Record Settlement", 
    "Compliance Tab → Settlement",
    "• For approved trades, click 'Record Settlement'\n• Confirm transfer details\n• Click 'Confirm Settlement'",
    "• Settlement confirmation dialog\n• Seller account → Buyer account transfer preview\n• Quantity and price summary\n• Digital signature required",
    "✓ Credits transferred in registry.registry_accounts\n✓ Seller balance decreased\n✓ Buyer balance increased\n✓ Settlement record created\n✓ Audit event: marketplace.settlement.recorded",
    "Settlement is irreversible; verify all details before confirming")

doc.add_page_break()

# =============================================================================
# PHASE 11: CREDIT RETIREMENT
# =============================================================================
doc.add_heading("PHASE 11: Credit Retirement", level=1)

add_ui_step(doc, 23, "Retire Credits for Compliance", 
    "Article 6 Ops Tab",
    "• Click 'Article 6 Ops' in sidebar\n• Select buyer account with owned credits\n• Select credits to retire\n• Enter retirement purpose (e.g., 'CORSIA 2025 compliance')\n• Enter beneficiary name\n• Click 'Retire Credits'",
    "• Portfolio view of owned credits\n• Multi-select credit rows\n• Retirement purpose text field\n• Beneficiary input\n• Retirement certificate preview",
    "✓ Credits marked as 'retired' in registry\n✓ Retirement certificate generated\n✓ Public disclosure record created\n✓ NDC accounting snapshot updated\n✓ Corresponding adjustment applied (if Article 6)\n✓ Audit event: carbon.credits.retired",
    "Retired credits are permanently non-transferable")

add_ui_step(doc, 24, "Verify Retirement Certificate", 
    "Article 6 Ops Tab → Certificate",
    "• Click 'Verify Certificate' button\n• Enter certificate ID (or use latest)\n• Review certificate details",
    "• Retirement certificate display:\n  - Certificate ID\n  - Serial numbers retired\n  - Quantity tCO2e\n  - Retirement date\n  - Beneficiary\n  - Purpose\n  - Verification hash\n• Public lookup link",
    "✓ Certificate publicly verifiable\n✓ Immutable record in audit trail\n✓ UNFCCC reporting ready",
    "Certificates can be verified by external parties without login")

doc.add_page_break()

# =============================================================================
# PHASE 12: NATIONAL REPORTING
# =============================================================================
doc.add_heading("PHASE 12: National Reporting & Accounting", level=1)

add_ui_step(doc, 25, "Generate National Accounting Snapshot", 
    "Reporting Tab",
    "• Click 'Reporting' in sidebar\n• Select reporting year\n• Select NDC sector\n• Click 'Lock Accounting Snapshot'",
    "• Reporting dashboard\n• Year and sector selectors\n• Current statistics:\n  - issued_tco2e\n  - retired_tco2e\n  - authorized_itmo_tco2e\n  - corresponding_adjustment_tco2e\n• 'Lock Snapshot' button",
    "✓ Snapshot saved to reporting.national_accounting_snapshots\n✓ data_hash computed for integrity\n✓ Snapshot locked (read-only)\n✓ Ready for BTR/ETF submission to UNFCCC",
    "Locked snapshots cannot be modified; create new snapshot if corrections needed")

add_ui_step(doc, 26, "Review National Deployment Maturity", 
    "National Stages Tab",
    "• Click 'National Stages' in sidebar\n• Review maturity score (0-100)\n• Review stage completion status\n• Execute any pending stage controls",
    "• National Deployment Control Board header\n• Maturity score display (large number)\n• Stage cards with:\n  - Stage name and number\n  - Status: not_started / in_progress / completed / blocked\n  - Required capabilities\n  - Current gaps\n  - Next controls\n• Execute Stage Controls buttons",
    "✓ National readiness tracked\n✓ Stage decisions recorded\n✓ Deployment blockers identified",
    "Target maturity score: 100 for full national deployment")

doc.add_page_break()

# =============================================================================
# SUMMARY
# =============================================================================
doc.add_heading("Process Flow Summary", level=1)

doc.add_paragraph("The complete admin-only project lifecycle consists of 26 UI steps across 12 phases:")

summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = "Table Grid"
headers = ["Phase", "Steps", "Key UI Tabs Used", "Outcome"]
for i, h in enumerate(headers):
    summary_table.rows[0].cells[i].text = h
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(summary_table.rows[0].cells[i], "0f766e")
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

phases = [
    ("1. Login & Setup", "1-2", "Login, Dashboard", "Admin authenticated, system healthy"),
    ("2. Project Registration", "3-5", "Carbon Registry", "New project in draft status"),
    ("3. Lifecycle Advancement", "6", "Project Lifecycle", "Project submitted for verification"),
    ("4. Evidence & Validation", "7-8", "Verification", "Evidence uploaded, auto-validated"),
    ("5. AI & GIS Reviews", "9-10", "AI Intelligence, GIS Intelligence", "AI and spatial assessments complete"),
    ("6. Regulatory Approval", "11-12", "Verification", "ZiCMA approves project"),
    ("7. Credit Issuance", "13-14", "Credit Registry", "Serialized credits issued"),
    ("8. Blockchain Anchoring", "15-18", "Anchoring", "Merkle root anchored, chain verified"),
    ("9. Marketplace", "19-20", "Marketplace", "Credits listed and traded"),
    ("10. Compliance & Settlement", "21-22", "Compliance", "Trade approved, credits transferred"),
    ("11. Retirement", "23-24", "Article 6 Ops", "Credits retired, certificate issued"),
    ("12. National Reporting", "25-26", "Reporting, National Stages", "Accounting locked, maturity tracked"),
]

for phase, steps, tabs, outcome in phases:
    row = summary_table.add_row().cells
    row[0].text = phase
    row[1].text = steps
    row[2].text = tabs
    row[3].text = outcome
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    "Document generated by ADMIN_UI_PROCESS_FLOW.docx.py  |  "
    "For API details, see docs/api/  |  "
    "For testing procedures, see DEPLOYMENT_AND_TESTING_GUIDE.docx"
)

# Save
doc.save("ADMIN_UI_PROCESS_FLOW.docx")
print("DOCX created: ADMIN_UI_PROCESS_FLOW.docx")
