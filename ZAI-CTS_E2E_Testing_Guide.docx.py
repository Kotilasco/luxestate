#!/usr/bin/env python3
"""Generate ZAI-CTS E2E Testing Guide as DOCX."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('ZAI-CTS End-to-End Testing Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.add_run('Zimbabwe AI-Enhanced Carbon Trading System').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Version 1.0 | 22 June 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Table of Contents placeholder
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. System Overview',
    '2. Service Endpoints',
    '3. Role-Based Testing Flows',
    '   3.1 ZiCMA Administrator (Super Admin)',
    '   3.2 Project Developer',
    '   3.3 Accredited Verifier',
    '   3.4 GIS Analyst',
    '   3.5 Buyer (Carbon Credit Purchaser)',
    '   3.6 Seller (Credit Owner)',
    '   3.7 Government Officer',
    '4. Complete Workflow: Project to Credit Retirement',
    '5. National Approval Workflow (Multi-Stage)',
    '6. Article 6 Workflow Tests',
    '7. Credit Lifecycle Tests',
    '8. GIS Validation Tests',
    '9. Fraud & Abuse Tests',
    '10. Security & RBAC Tests',
    '11. Evidence Validation Tests',
    '12. Marketplace Edge Case Tests',
    '13. National Reporting Tests',
    '14. Performance & Load Tests',
    '15. Blockchain Integration Tests',
    '16. API Contract & Integration Tests',
    '17. Notification & Alert Tests',
    '18. Mobile & Responsive Tests',
    '19. Accessibility (WCAG) Tests',
    '20. Disaster Recovery & Backup Tests',
    '21. Session & Authentication Security',
    '22. Data Integrity & Migration Tests',
    '23. Search & Export Functionality',
    '24. Troubleshooting'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. System Overview
doc.add_heading('1. System Overview', level=1)
doc.add_paragraph(
    'The ZAI-CTS (Zimbabwe AI-Enhanced Carbon Trading System) is a national carbon registry '
    'platform that manages the full lifecycle of carbon projects—from registration and verification '
    'to credit issuance, marketplace trading, and retirement. The platform supports multiple user '
    'roles with role-based access control (RBAC).'
)

# 2. Service Endpoints
doc.add_heading('2. Service Endpoints', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Service'
hdr_cells[1].text = 'Port'
hdr_cells[2].text = 'Health Check URL'
services = [
    ('Next.js Frontend', '3000', 'http://localhost:3000'),
    ('API Gateway', '8082', 'http://127.0.0.1:8082/health'),
    ('Carbon Registry', '8102', 'http://127.0.0.1:8102/health'),
    ('AI Validation', '8103', 'http://127.0.0.1:8103/health'),
    ('GIS Service', '8104', 'http://127.0.0.1:8104/health'),
    ('Marketplace', '8106', 'http://127.0.0.1:8106/health'),
    ('Compliance', '8107', 'http://127.0.0.1:8107/health'),
]
for svc, port, url in services:
    row_cells = table.add_row().cells
    row_cells[0].text = svc
    row_cells[1].text = port
    row_cells[2].text = url

doc.add_paragraph()

# 3. Role-Based Testing Flows
doc.add_heading('3. Role-Based Testing Flows', level=1)

def add_role_section(doc, role, credentials, tasks):
    doc.add_heading(role, level=2)
    if credentials:
        p = doc.add_paragraph()
        p.add_run('Login Credentials: ').bold = True
        p.add_run(credentials)
    for i, (task, steps, expected) in enumerate(tasks, 1):
        doc.add_heading(f'{i}. {task}', level=3)
        p = doc.add_paragraph()
        p.add_run('Steps:\n').bold = True
        for step in steps:
            doc.add_paragraph(step, style='List Bullet')
        p = doc.add_paragraph()
        p.add_run('Expected Result: ').bold = True
        p.add_run(expected)
        doc.add_paragraph()

# 3.1 ZiCMA Administrator
add_role_section(doc, '3.1 ZiCMA Administrator (Super Admin)',
    'admin@zai-cts.gov.zw / ZAI-CTS-Admin-2026!',
    [
        ('View Dashboard',
         ['Navigate to http://localhost:3000', 'Enter admin credentials', 'Click Login'],
         'Dashboard loads with stats: 5 Projects, 0 Credits, 0 Evidence, Gateway: healthy'),
        ('Browse Registered Projects',
         ['Click "Carbon Registry" in sidebar'],
         'List shows 5 projects: Kariba AI Verified Forest, Hwange Solar Energy, Chipinge Agroforestry, Mutare Reforestation'),
        ('Submit Project for Verification',
         ['In Carbon Registry, select a project', 'Click "Submit for Verification" button'],
         'Project status changes from draft to submitted. Alert confirms submission.'),
        ('Run Full Verification Pipeline',
         ['Click "Verification" tab', 'Click "Start Case"', 'Run steps: Evidence → Automatic → AI → GIS → MRV → Verifier → ZiCMA',
          'For each step, click the step button and confirm'],
         'Verification case progresses through all stages. Final ZiCMA approval changes status to approved.'),
        ('Issue Carbon Credits',
         ['Click "Credit Registry" tab', 'Select approved project', 'Click "Issue Credits"',
          'Enter vintage year and quantity', 'Confirm issuance'],
         'Credits issued with unique serial prefix. Blockchain transaction ID recorded. Status: issued.'),
        ('Review GIS Intelligence',
         ['Click "GIS Intelligence" tab', 'Select project with GIS data'],
         '6 layers visible: Project boundary, Sentinel-2 composite, Forest cover, NASA FIRMS fire alerts, Carbon density, MRV field plots. Map inspector shows Lat/Lng/Area.'),
        ('Check Marketplace Pricing',
         ['Click "Marketplace" tab', 'View "Dynamic Pricing" sub-tab'],
         'Real-time ITMO prices displayed by vintage year ($27–$42/tCO2e). Demand and authorization percentages shown.'),
        ('Manage Users & Audit Trail',
         ['Click "Identity & Users" tab', 'Review pending approvals', 'Click "Reporting" tab'],
         'User list with roles visible. Audit events show all actions with actor, timestamp, and outcome.'),
    ])

doc.add_page_break()

# 3.2 Project Developer
add_role_section(doc, '3.2 Project Developer',
    'Register first, then use approved credentials',
    [
        ('Register New Account',
         ['On login screen, click "Register" tab',
          'Fill: Full name (e.g., John Mutasa), Email (developer@greenzim.org), Password (DevPass2026!)',
          'Select Role: "Project Developer"',
          'Fill Organization: GreenZim Carbon, Type: Carbon Developer, Registration: PVT2026-001',
          'Click "Create account"'],
         'Success message: "John Mutasa registered. A registry administrator must approve the account before login."'),
        ('Login After Approval',
         ['Wait for admin approval', 'Enter credentials', 'Click Login'],
         'Dashboard loads with Project Developer role. Sidebar shows relevant tabs.'),
        ('Register a New Carbon Project',
         ['Click "Carbon Registry" tab', 'Fill Project code (e.g., RE-20260001)',
          'Fill Title, Description, Methodology (e.g., VM0047 Afforestation Reforestation Revegetation)',
          'Select Province and District', 'Enter Estimated annual tCO2e (e.g., 125000)',
          'Set Start date and Crediting years (e.g., 30)',
          'Click "Register Project"'],
         'Project appears in registered projects list with status: draft. Audit event recorded.'),
        ('Submit Project for Verification',
         ['Select the newly registered project', 'Click "Submit for Verification"'],
         'Project status changes to submitted. Alert confirms. Awaiting verifier assignment.'),
        ('Upload Evidence Package',
         ['Click "Project Lifecycle" tab', 'Drag and drop ZIP file or individual documents',
          'Documents auto-categorized: PDD, Feasibility study, Environmental impact, GIS boundary, Satellite imagery, Photos'],
         'Evidence package uploaded. System stores file hashes. Upload progress shown.'),
    ])

doc.add_page_break()

# 3.3 Accredited Verifier
add_role_section(doc, '3.3 Accredited Verifier',
    'Register with role "Accredited Verifier", await admin approval',
    [
        ('Review Evidence Package',
         ['Login as Accredited Verifier', 'Click "Verification" tab',
          'Select project with active verification case', 'Review uploaded evidence documents'],
         'Evidence package visible with all categories. File hashes and upload timestamps shown.'),
        ('Run Verification Steps',
         ['Click "Start Case" (if not started)', 'Run through each verification stage:',
          '  • Automatic validation: System runs automated checks',
          '  • AI review: AI validates ownership, risk, and integrity scores',
          '  • GIS review: Spatial analysis and satellite imagery validation',
          '  • MRV review: Field measurement and monitoring data review'],
         'Each step updates status from pending → in_progress → complete. Scores update.'),
        ('Verifier Approval with Digital Signature',
         ['Review all completed steps', 'Click "Verifier Review" button',
          'Confirm digital signature', 'Add verifier notes'],
         'Status changes to verifier_approved. Digital signature hash recorded. Audit trail updated.'),
        ('Request More Information',
         ['If evidence insufficient, click "Request More Information"',
          'Select required document categories', 'Add detailed notes'],
         'Project status changes to revision_requested. Developer notified. Case on hold.'),
    ])

doc.add_page_break()

# 3.4 GIS Analyst
add_role_section(doc, '3.4 GIS Analyst',
    'Register with role "GIS Analyst", await admin approval',
    [
        ('Run GIS Assessment',
         ['Login as GIS Analyst', 'Click "GIS Intelligence" tab',
          'Select project from dropdown', 'Click "Run GIS Assessment"'],
         'Assessment initiated. Progress bars show for each layer.'),
        ('Review Layer Stack',
         ['Check each layer status in sidebar:',
          '  • Project boundary: Vector - validated',
          '  • Sentinel-2 composite: Satellite - ready',
          '  • Forest cover: Raster - analyzed',
          '  • NASA FIRMS fire alerts: Risk - monitored',
          '  • Carbon density: Model - computed',
          '  • MRV field plots: Survey - validated'],
         'All layers show green/completed status. Map inspector populated with coordinates.'),
        ('Submit GIS Evidence',
         ['In Evidence Intake section, verify:', '  • Boundary GeoJSON (polygon coordinates)',
          '  • Satellite scene ID (e.g., S2A_MSIL2A_20260615T073621_KARIBA)',
          '  • Land-cover source (ESA WorldCover 10m)',
          '  • Fire-alert source (NASA FIRMS VIIRS)',
          '  • Field MRV reference',
          'Click "Submit Evidence"'],
         'Evidence submitted. Spatial findings generated. Layer health updated to validated.'),
        ('Switch Basemaps',
         ['Click basemap buttons: Imagery, Terrain, Land cover, Hybrid'],
         'Map updates with selected basemap. All 4 modes functional.'),
    ])

doc.add_page_break()

# 3.5 Buyer
add_role_section(doc, '3.5 Buyer (Carbon Credit Purchaser)',
    'Register with role "Buyer", await admin approval',
    [
        ('Browse Marketplace Pricing',
         ['Login as Buyer', 'Click "Marketplace" tab',
          'View "Dynamic Pricing" sub-tab'],
         'Current market prices shown by vintage: 2020 ($27–$33), 2021 ($32–$36), 2022 ($37–$40), 2023 ($31–$42). Demand % and auth % displayed.'),
        ('Browse Credit Listings',
         ['Click "Browse Listings" sub-tab'],
         'Available credit batches listed with serial prefix, vintage, quantity, price, seller, authorization status.'),
        ('Calculate Purchase Price',
         ['In "Dynamic Pricing" tab, locate Price Calculator',
          'Enter Quantity in tCO2e (e.g., 1000)',
          'Click "Calculate Price"'],
         'Estimated total price displayed based on current market rate and selected vintage.'),
        ('Purchase Credits',
         ['Select a listing', 'Click "Buy" or "Make Offer"',
          'Confirm purchase quantity and total',
          'Complete transaction'],
         'Credits transferred to buyer account. Transaction recorded in Compliance → Retire/Transfer. Marketplace listing updated.'),
    ])

doc.add_page_break()

# 3.6 Seller
add_role_section(doc, '3.6 Seller (Credit Owner)',
    'Register with role "Seller", await admin approval',
    [
        ('Create Marketplace Listing',
         ['Login as Seller', 'Click "Marketplace" tab',
          'Click "Create Listing" or "Sell Credits"',
          'Select credit batch from inventory',
          'Set price per tCO2e',
          'Set vintage year',
          'Set authorization status (Authorized = ITMO eligible)',
          'Enter total quantity to list'],
         'Listing created. Appears in "Browse Listings" for buyers. Status: active.'),
        ('Manage Active Listings',
         ['View "My Listings" section'],
         'All seller listings shown with status: active, sold, expired. Quantities and prices editable.'),
        ('View Sales History',
         ['Click "Compliance" tab → "Transfer History"'],
         'All sales transactions listed with buyer, quantity, price, date, transaction hash.'),
    ])

doc.add_page_break()

# 3.7 Government Officer
add_role_section(doc, '3.7 Government Officer',
    'Register with role "Government Officer", await admin approval',
    [
        ('Review Compliance & Retirements',
         ['Login as Government Officer', 'Click "Compliance" tab'],
         'All retirement records visible. Filter by date range, project, beneficiary, purpose.'),
        ('Generate National Reports',
         ['Click "Reporting" tab', 'Select report type: National Inventory, NDC Tracking, or Article 6 Report',
          'Set date range', 'Click "Generate Report"'],
         'Report generated with charts and tables. Exportable to PDF/Excel.'),
        ('Monitor Verification Pipeline',
         ['Click "National Stages" tab'],
         'Overview of all projects by stage: draft, submitted, verification, approved, issued.'),
        ('Review Audit Logs',
         ['Click "Administration" tab', 'View "Audit Logs"'],
         'Complete audit trail: actor, role, action, outcome, timestamp, correlation ID. Searchable and filterable.'),
    ])

doc.add_page_break()

# 4. Complete Workflow
doc.add_heading('4. Complete Workflow: Project Registration to Credit Retirement', level=1)
workflow_steps = [
    ('1. Project Developer', 'Registers organization and project', 'Project status: draft'),
    ('2. Project Developer', 'Submits project for verification', 'Project status: submitted'),
    ('3. ZiCMA Administrator', 'Assigns verifier, starts verification case', 'Case status: started'),
    ('4. Automatic System', 'Runs automated validation checks', 'Step: automatic - complete'),
    ('5. AI Validation Service', 'AI reviews project data and evidence', 'Step: ai - complete. Scores: integrity, confidence, risk'),
    ('6. GIS Analyst', 'Runs spatial assessment and layer analysis', 'Step: gis - complete. All layers validated'),
    ('7. MRV Officer', 'Reviews field measurements and monitoring', 'Step: mrv - complete'),
    ('8. Accredited Verifier', 'Reviews all evidence, digitally signs approval', 'Step: verifier - complete. Status: verifier_approved'),
    ('9. ZiCMA Administrator', 'Final regulatory approval', 'Step: zicma - complete. Project status: approved'),
    ('10. ZiCMA Administrator', 'Issues carbon credit batch', 'Credits issued with serial numbers. Status: issued'),
    ('11. Seller', 'Lists credits on marketplace', 'Listing active. Buyers can view and purchase'),
    ('12. Buyer', 'Purchases credits', 'Credits transferred to buyer account'),
    ('13. Buyer / Government Officer', 'Retires credits for compliance', 'Credits retired. Retirement certificate generated. Blockchain recorded.'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Actor'
hdr[1].text = 'Action'
hdr[2].text = 'Result'
for actor, action, result in workflow_steps:
    row = table.add_row().cells
    row[0].text = actor
    row[1].text = action
    row[2].text = result

doc.add_page_break()

# 5. National Approval Workflow (Multi-Stage)
doc.add_heading('5. National Approval Workflow (Multi-Stage)', level=1)
doc.add_paragraph(
    'The complete national approval pipeline for carbon projects in Zimbabwe includes multiple '
    'screening stages before a project reaches ZiCMA final approval. Each stage is a gate that '
    'must be passed.'
)
national_workflow = [
    ('1. Developer Submission', 'Project Developer', 'Project registered and submitted for verification.'),
    ('2. Technical Screening', 'Technical Reviewer', 'Methodology validity, data completeness, baseline plausibility checked.'),
    ('3. GIS Screening', 'GIS Analyst', 'Boundary overlap with protected areas, mining concessions, national parks checked. Leakage and permanence risks assessed.'),
    ('4. Environmental Review', 'Environmental Officer', 'Environmental impact assessment reviewed. No objection certificate verified.'),
    ('5. Accredited Verifier Review', 'Accredited Verifier', 'Independent verification of project design, monitoring plan, and evidence.'),
    ('6. National Approval Committee', 'Committee Members', 'Multi-stakeholder review: EMA, Forestry Commission, ZiCMA, Local Authority.'),
    ('7. ZiCMA Final Approval', 'ZiCMA Administrator', 'Regulatory sign-off. Project status: approved.'),
    ('8. Credit Issuance', 'ZiCMA Administrator', 'Serialized credits issued with unique prefix. Blockchain recorded.'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Stage'
hdr[1].text = 'Actor'
hdr[2].text = 'Outcome'
for stage, actor, outcome in national_workflow:
    row = table.add_row().cells
    row[0].text = stage
    row[1].text = actor
    row[2].text = outcome

doc.add_paragraph()
doc.add_heading('5.1 Technical Screening Tests', level=2)
tech_tests = [
    ('Methodology not on approved list', 'Submit VM9999 → Expected: Rejected at Technical Screening'),
    ('Missing mandatory PDD section', 'Upload incomplete PDD → Expected: Revision requested'),
    ('Baseline exceeds historical data', 'Baseline tCO2e > 150% of regional average → Expected: Flagged for review'),
    ('Crediting period > 30 years', 'Enter 35 years → Expected: Rejected (max 30)'),
    ('Negative estimated tCO2e', 'Enter -500 → Expected: Validation error (must be ≥ 0)'),
]
for test, expected in tech_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('5.2 GIS Screening Tests', level=2)
gis_screen_tests = [
    ('Overlap with national park', 'Boundary intersects Hwange National Park → Expected: Rejected'),
    ('Overlap with existing project', 'Boundary overlaps RE-20240001 → Expected: Rejected with conflict report'),
    ('Overlap with mining concession', 'Boundary intersects mining lease → Expected: Rejected, referral to Mines Ministry'),
    ('Buffer zone violation', 'Project < 500m from protected area boundary → Expected: Flagged, requires EMA waiver'),
]
for test, expected in gis_screen_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('5.3 Environmental Review Tests', level=2)
env_tests = [
    ('No EIA certificate', 'Project > 100 hectares without EIA → Expected: Rejected, EIA required'),
    ('Expired EIA', 'EIA older than 5 years → Expected: Rejected, EIA renewal required'),
    ('Water authority objection', 'ZINWA raises objection → Expected: On hold pending resolution'),
    ('Forestry Commission objection', 'Forestry raises objection → Expected: On hold, mediation required'),
]
for test, expected in env_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('5.4 National Approval Committee Tests', level=2)
comm_tests = [
    ('Committee quorum not met', '< 4 of 7 members present → Expected: Vote postponed'),
    ('Tie vote', '3 approve, 3 reject, 1 abstain → Expected: Chair casting vote required'),
    ('Conditional approval', 'Committee approves with conditions → Expected: Status: conditionally_approved'),
    ('Rejection with reasons', 'Committee rejects → Expected: Detailed rejection letter with appeal rights'),
]
for test, expected in comm_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 6. Article 6 Workflow Tests
doc.add_heading('6. Article 6 Workflow Tests', level=1)
doc.add_paragraph(
    'Article 6 of the Paris Agreement governs international carbon markets. '
    'These tests verify compliance with Article 6.2 (bilateral/cooperative approaches) '
    'and Article 6.4 (mechanism).'
)
doc.add_heading('6.1 Host Country Authorization', level=2)
art6_host = [
    ('Request host country authorization', 'ZiCMA submits authorization request to UNFCCC registry'),
    ('Authorization pending', 'Status: pending_authorization. Credits locked.'),
    ('Authorization granted', 'UNFCCC issues authorization ID. Status: authorized.'),
    ('Authorization denied', 'UNFCCC denies. Status: authorization_denied. Credits remain domestic.'),
]
for step, outcome in art6_host:
    p = doc.add_paragraph()
    p.add_run(f'{step}: ').bold = True
    p.add_run(outcome)

doc.add_heading('6.2 Corresponding Adjustments', level=2)
art6_ca = [
    ('Calculate corresponding adjustment', 'Based on NDC target and transferred credits'),
    ('First transfer adjustment', 'Transfer 1000 tCO2e → NDC subtracts 1000 from achievement'),
    ('Subsequent transfer adjustment', 'Transfer another 500 → Cumulative adjustment: 1500'),
    ('Reversal adjustment', 'If credits reversed → NDC adds back reversed amount'),
]
for step, outcome in art6_ca:
    p = doc.add_paragraph()
    p.add_run(f'{step}: ').bold = True
    p.add_run(outcome)

doc.add_heading('6.3 International Transfer Authorization', level=2)
art6_transfer = [
    ('Buyer requests ITMO transfer', 'Buyer from Switzerland requests 500 tCO2e'),
    ('Transfer authorization check', 'System verifies: authorization ID valid, not expired, not revoked'),
    ('Corresponding adjustment applied', 'Zimbabwe NDC adjusted by 500 tCO2e'),
    ('Transfer executed', 'Credits moved to Swiss registry. Blockchain TX recorded.'),
    ('Double transfer prevention', 'Attempt same transfer again → Expected: Rejected (already transferred)'),
]
for step, outcome in art6_transfer:
    p = doc.add_paragraph()
    p.add_run(f'{step}: ').bold = True
    p.add_run(outcome)

doc.add_heading('6.4 Article 6 Registry Export', level=2)
art6_export = [
    ('Generate Article 6 export package', 'Include: project details, authorization ID, transfer history, corresponding adjustments'),
    ('Export to Swiss registry', 'Package sent via UNFCCC international registry'),
    ('Export validation', 'Swiss registry validates package integrity and signatures'),
    ('Import confirmation', 'Swiss registry confirms import. Status: exported.'),
]
for step, outcome in art6_export:
    p = doc.add_paragraph()
    p.add_run(f'{step}: ').bold = True
    p.add_run(outcome)

doc.add_heading('6.5 Double Counting Prevention Tests', level=2)
art6_double = [
    ('Same credit in two registries', 'Attempt to register same serial in Swiss and Swedish registries → Expected: Second registry rejects'),
    ('Same credit retired twice', 'Retire in Zimbabwe registry, attempt retire in buyer registry → Expected: Rejected'),
    ('Same credit transferred twice', 'Transfer to Buyer A, attempt transfer same serial to Buyer B → Expected: Rejected'),
    ('NDC and voluntary market', 'Credit counted in NDC, attempt voluntary market sale → Expected: Flagged, requires adjustment clarification'),
]
for test, expected in art6_double:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 7. Credit Lifecycle Tests
doc.add_heading('7. Credit Lifecycle Tests', level=1)
doc.add_paragraph('Comprehensive testing of all credit states from issuance to final disposition.')

doc.add_heading('7.1 Transfer Chain', level=2)
doc.add_paragraph('Issue → Transfer → Transfer → Transfer → Retire')
transfer_chain = [
    ('1. Issue', 'ZiCMA issues 1000 credits (Serial: ZWE-RE-2025-00001 to 01000)'),
    ('2. First Transfer', 'Transfer 1000 from ZiCMA to Seller A → Seller A owns 1000'),
    ('3. Second Transfer', 'Seller A sells 600 to Buyer B → Seller A: 400, Buyer B: 600'),
    ('4. Third Transfer', 'Buyer B sells 200 to Buyer C → Buyer B: 400, Buyer C: 200'),
    ('5. Retire', 'Buyer C retires 200 → Buyer C: 0. Status: retired. Certificate generated.'),
]
for step, outcome in transfer_chain:
    p = doc.add_paragraph()
    p.add_run(f'{step}: ').bold = True
    p.add_run(outcome)
doc.add_paragraph('At each step: verify ownership, serial numbers, audit trail, blockchain record.')

doc.add_heading('7.2 Cancellation', level=2)
cancel_tests = [
    ('Cancel before any transfer', 'Issue 500 → Cancel 500 → Status: cancelled. No owner.'),
    ('Cancel partial holding', 'Seller has 400, cancels 100 → Seller: 300. Cancelled: 100.'),
    ('Cancel by non-owner', 'Buyer attempts to cancel Seller credits → Expected: Rejected (not owner)'),
    ('Cancel already retired', 'Attempt to cancel retired credits → Expected: Rejected (final state)'),
]
for test, expected in cancel_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('7.3 Reversal / Revocation', level=2)
revoke_tests = [
    ('Revoke for non-compliance', 'Project found non-compliant → ZiCMA revokes 1000 credits'),
    ('Revoke partially', 'Revoke 300 of 1000 → 700 remain active. 300 revoked.'),
    ('Revoke already transferred', 'Credits transferred to Buyer → ZiCMA revokes → Buyer notified. Chain traced.'),
    ('Revoke already retired', 'Attempt revoke retired credits → Expected: Rejected (final state)'),
    ('Revoke without authority', 'Seller attempts to revoke own credits → Expected: Rejected (only ZiCMA can revoke)'),
]
for test, expected in revoke_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('7.4 Buffer Pool', level=2)
buffer_tests = [
    ('Allocate to buffer on issuance', 'Issue 1000, allocate 200 (20%) to buffer pool → Active: 800, Buffer: 200'),
    ('Release from buffer', 'After 5 years with no reversal, release 100 from buffer → Active: 900, Buffer: 100'),
    ('Use buffer for reversal', 'Project reversal of 150 → Draw 150 from buffer → Buffer: 50. Active unaffected.'),
    ('Buffer depletion', 'Reversal of 250, buffer only 200 → Draw 200 from buffer, revoke 50 from active'),
]
for test, expected in buffer_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('7.5 Credit State Machine', level=2)
doc.add_paragraph('Valid state transitions:')
state_transitions = [
    'issued → transferred → transferred → retired (valid)',
    'issued → transferred → cancelled (valid)',
    'issued → revoked (valid)',
    'issued → retired (valid, direct retirement by issuer)',
    'retired → transferred (INVALID: final state)',
    'cancelled → transferred (INVALID: final state)',
    'revoked → retired (INVALID: final state)',
]
for transition in state_transitions:
    doc.add_paragraph(transition, style='List Bullet')

doc.add_page_break()

# 8. GIS Validation Tests
doc.add_heading('8. GIS Validation Tests', level=1)
doc.add_paragraph('Comprehensive spatial validation for carbon project boundaries.')

doc.add_heading('8.1 Boundary Overlap Tests', level=2)
overlap_tests = [
    ('Protected area overlap', 'Boundary intersects Victoria Falls National Park → Expected: Rejected. Conflict report generated.'),
    ('Existing carbon project overlap', 'Boundary intersects RE-20240001 (Mutare Reforestation) → Expected: Rejected with overlap percentage.'),
    ('Mining concession overlap', 'Boundary intersects registered mining lease → Expected: Rejected. Referral to Ministry of Mines.'),
    ('Community land overlap', 'Boundary intersects registered communal land without consent → Expected: Flagged. Consent required.'),
    ('Cross-border overlap', 'Boundary crosses into Zambia → Expected: Rejected. International boundary violation.'),
]
for test, expected in overlap_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('8.2 Leakage Analysis Tests', level=2)
leakage_tests = [
    ('Deforestation leakage', 'Satellite shows increased deforestation within 5km buffer of project boundary → Expected: Leakage detected. Quantify tCO2e.'),
    ('Agricultural leakage', 'Cropland expansion within 2km buffer → Expected: Flagged for MRV review.'),
    ('Urban expansion leakage', 'Settlement expansion within buffer → Expected: Documented in leakage report.'),
    ('No leakage', 'No land-use change within 5km buffer → Expected: Leakage analysis passed.'),
]
for test, expected in leakage_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('8.3 Permanence Risk Tests', level=2)
perm_tests = [
    ('Fire risk assessment', 'NASA FIRMS shows 3+ fire alerts within boundary in past 12 months → Expected: High fire risk. Buffer pool increased.'),
    ('Flood risk assessment', 'Historical flood data shows recurring flooding → Expected: Flood risk documented. Mitigation plan required.'),
    ('Drought risk assessment', 'SPI index < -1.5 for 2+ years → Expected: Drought risk flagged. Monitoring frequency increased.'),
    ('Political risk assessment', 'Project in disputed area → Expected: Political risk flagged. Insurance/buffer requirements.'),
    ('Low permanence risk', 'No fire/flood/drought/political risks → Expected: Standard buffer allocation.'),
]
for test, expected in perm_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('8.4 Additionality Tests', level=2)
add_tests = [
    ('Historical land use', 'Area was primary forest in 2000 → Expected: Not additional (would have remained forest).'),
    ('Baseline deforestation rate', 'Regional deforestation 2%/year, project area stable → Expected: Additional (prevented deforestation).'),
    ('Financial additionality', 'Project IRR < 8% without carbon revenue → Expected: Financially additional.'),
    ('Regulatory additionality', 'Project exceeds legal requirements → Expected: Regulatory additionality confirmed.'),
    ('Common practice', 'No similar projects within 50km → Expected: Common practice test passed.'),
]
for test, expected in add_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('8.5 GIS Evidence Integrity Tests', level=2)
gis_ev_tests = [
    ('Corrupted GeoJSON', 'Upload malformed GeoJSON → Expected: Rejected. Parse error.'),
    ('Wrong CRS', 'GeoJSON in EPSG:3857 instead of EPSG:4326 → Expected: Rejected. CRS mismatch.'),
    ('Area mismatch', 'GeoJSON area = 500ha, claimed area = 1000ha → Expected: Flagged. Area discrepancy.'),
    ('Self-intersecting polygon', 'Polygon boundary crosses itself → Expected: Rejected. Invalid geometry.'),
    ('Missing boundary', 'No boundary file uploaded → Expected: Rejected. Mandatory evidence missing.'),
]
for test, expected in gis_ev_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 9. Fraud & Abuse Tests
doc.add_heading('9. Fraud & Abuse Tests', level=1)
doc.add_paragraph('Deliberate attempts to compromise registry integrity and expected system responses.')

fraud_tests = [
    ('Duplicate project submission', 'Submit same polygon with different project code → Expected: Rejected. Duplicate boundary detected.'),
    ('Fake satellite imagery', 'Upload Sentinel-2 image from Kariba district for Mutare project → Expected: GIS mismatch detected. Metadata mismatch.'),
    ('Fake boundary (another country)', 'Upload GeoJSON from Kenya → Expected: Rejected. Outside Zimbabwe.'),
    ('Double issuance attempt', 'Issue credits for RE-20250001, then attempt second issuance for same project/vintage → Expected: Blocked. Already issued.'),
    ('Double retirement attempt', 'Retire serial ZWE-RE-2025-00001, attempt retire same serial again → Expected: Rejected. Already retired.'),
    ('Transfer non-existent credits', 'Attempt transfer serial ZWE-RE-2099-99999 → Expected: Rejected. Serial not found.'),
    ('Transfer without ownership', 'Buyer B attempts to transfer Seller A credits → Expected: Rejected. Not owner.'),
    ('Backdated project', 'Set start date to 2010 (before platform launch) → Expected: Rejected. Invalid date.'),
    ('Future-dated project', 'Set start date to 2035 → Expected: Rejected. Start date cannot be > 2 years future.'),
    ('Inflated baseline', 'Claim 1,000,000 tCO2e/year for 10-hectare forest → Expected: Flagged. Exceeds theoretical maximum.'),
    ('Fake verifier signature', 'Submit verification without valid accredited verifier → Expected: Rejected. Verifier not accredited.'),
    ('Expired verifier approval', 'Verifier accreditation expired 2025, signs in 2026 → Expected: Rejected. Verifier not active.'),
    ('SQL injection in project code', 'Enter project code: RE-2025\'; DROP TABLE-- → Expected: Rejected. Input sanitized.'),
    ('XSS in project title', 'Enter title: <script>alert(1)</script> → Expected: Sanitized. Script not executed.'),
    ('Brute force login', '100 failed login attempts in 1 minute → Expected: Account locked for 30 minutes.'),
    ('Session hijacking', 'Attempt use stolen JWT token after logout → Expected: Rejected. Token invalidated.'),
    ('Rate limit bypass', 'Submit 1000 API requests/second → Expected: Rate limited (429 Too Many Requests).'),
]
for test, expected in fraud_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 10. Security & RBAC Tests
doc.add_heading('10. Security & RBAC Tests', level=1)
doc.add_paragraph('Verify that users cannot perform actions outside their role permissions.')

doc.add_heading('10.1 Developer Restrictions', level=2)
dev_restrict = [
    ('Issue credits', 'Developer clicks Issue Credits → Expected: Rejected (403 Forbidden). Only ZiCMA can issue.'),
    ('Approve project', 'Developer attempts ZiCMA approval → Expected: Rejected.'),
    ('Retire credits', 'Developer attempts retirement → Expected: Rejected.'),
    ('Modify issued credits', 'Developer attempts edit credit batch → Expected: Rejected.'),
    ('Approve user registration', 'Developer attempts approve another user → Expected: Rejected.'),
]
for test, expected in dev_restrict:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('10.2 Verifier Restrictions', level=2)
ver_restrict = [
    ('Modify issued credits', 'Verifier attempts to change credit quantity → Expected: Rejected.'),
    ('Issue credits', 'Verifier attempts credit issuance → Expected: Rejected.'),
    ('Approve own verification', 'Verifier attempts to bypass committee → Expected: Rejected.'),
    ('Delete evidence', 'Verifier attempts to delete uploaded evidence → Expected: Rejected (immutable).'),
]
for test, expected in ver_restrict:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('10.3 Buyer Restrictions', level=2)
buy_restrict = [
    ('Change registry data', 'Buyer attempts to edit project details → Expected: Rejected.'),
    ('Issue credits', 'Buyer attempts credit issuance → Expected: Rejected.'),
    ('Approve project', 'Buyer attempts project approval → Expected: Rejected.'),
    ('View other buyer data', 'Buyer A attempts to view Buyer B purchases → Expected: Rejected.'),
]
for test, expected in buy_restrict:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('10.4 Government Officer Restrictions', level=2)
gov_restrict = [
    ('Edit evidence', 'Gov officer attempts to modify project evidence → Expected: Rejected.'),
    ('Issue credits', 'Gov officer attempts credit issuance → Expected: Rejected.'),
    ('Retire credits', 'Gov officer attempts retirement → Expected: Allowed (compliance role).'),
    ('Modify audit log', 'Gov officer attempts to delete audit event → Expected: Rejected (immutable).'),
]
for test, expected in gov_restrict:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('10.5 GIS Analyst Restrictions', level=2)
gis_restrict = [
    ('Modify project data', 'GIS Analyst attempts to edit project title → Expected: Rejected.'),
    ('Approve project', 'GIS Analyst attempts approval → Expected: Rejected.'),
    ('Run GIS assessment', 'GIS Analyst runs assessment → Expected: Allowed.'),
    ('Submit GIS evidence', 'GIS Analyst submits evidence → Expected: Allowed.'),
]
for test, expected in gis_restrict:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 11. Evidence Validation Tests
doc.add_heading('11. Evidence Validation Tests', level=1)
doc.add_paragraph('Tests for evidence package integrity, format, and completeness.')

evidence_tests = [
    ('Corrupted file', 'Upload corrupted ZIP (bad CRC) → Expected: Rejected. Integrity check failed.'),
    ('Wrong format', 'Upload .exe file as evidence → Expected: Rejected. File type not allowed.'),
    ('Duplicate file', 'Upload same PDF twice → Expected: Second upload flagged. Hash matches existing.'),
    ('Missing mandatory evidence', 'Submit without PDD → Expected: Rejected. PDD is mandatory.'),
    ('Missing verifier statement', 'Submit verification without verifier PDF → Expected: Rejected.'),
    ('Invalid signature', 'PDF signature certificate expired → Expected: Flagged. Signature invalid.'),
    ('Hash mismatch', 'File modified after upload, hash no longer matches → Expected: Flagged. Integrity compromised.'),
    ('File too large', 'Upload 500MB single file → Expected: Rejected. Exceeds 100MB limit.'),
    ('Empty file', 'Upload 0-byte PDF → Expected: Rejected. Empty file.'),
    ('Virus infected file', 'Upload file with EICAR test signature → Expected: Rejected. Virus detected.'),
    ('Wrong document category', 'Upload PDD as satellite imagery → Expected: Flagged. Category mismatch.'),
    ('Expired document', 'Upload EIA from 2010 (project starts 2026) → Expected: Flagged. Document expired.'),
    ('Unsigned document', 'Upload verifier statement without digital signature → Expected: Flagged. Signature required.'),
    ('Watermark tampering', 'Upload document with altered watermark → Expected: Flagged. Document authenticity failed.'),
]
for test, expected in evidence_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 12. Marketplace Edge Case Tests
doc.add_heading('12. Marketplace Edge Case Tests', level=1)
doc.add_paragraph('Boundary cases and concurrent transaction tests for the carbon credit marketplace.')

market_tests = [
    ('Partial purchase', '1000 credits listed. Buyer purchases 200 → Expected: Listing updated to 800 remaining.'),
    ('Multiple buyers', 'Buyer A buys 500, Buyer B buys 300 → Expected: Listing shows 200 remaining. Both transactions recorded.'),
    ('Oversell protection', '1000 listed. Attempt buy 1200 → Expected: Rejected. Insufficient quantity.'),
    ('Exact purchase', '1000 listed. Buy exactly 1000 → Expected: Listing status: sold. Removed from browse.'),
    ('Zero purchase', 'Attempt buy 0 → Expected: Rejected. Minimum quantity: 1.'),
    ('Fractional purchase', 'Attempt buy 0.5 tCO2e → Expected: Rejected. Whole credits only.'),
    ('Negative purchase', 'Attempt buy -100 → Expected: Rejected. Negative quantity.'),
    ('Price change mid-purchase', 'Seller changes price while Buyer checking out → Expected: Transaction uses price at time of order.'),
    ('Concurrent purchase', 'Buyer A and Buyer B simultaneously buy 600 each from 1000 listing → Expected: First completes (600), second gets 400 (partial) or rejected if atomic.'),
    ('Expired listing', 'Listing expired 30 days ago → Expected: Not visible in browse. Cannot purchase.'),
    ('Cancelled listing', 'Seller cancels listing before purchase → Expected: Not visible. Existing offers rejected.'),
    ('Buy own listing', 'Seller attempts to buy their own listing → Expected: Rejected. Self-dealing prohibited.'),
    ('Unapproved credits', 'Credits in draft status listed → Expected: Rejected. Only issued credits listable.'),
    ('Revoked credits listed', 'Revoked credits appear in listing → Expected: Rejected. Revoked credits cannot be traded.'),
]
for test, expected in market_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 13. National Reporting Tests
doc.add_heading('13. National Reporting Tests', level=1)
doc.add_paragraph('Comprehensive reporting for national carbon accounting and international transparency.')

doc.add_heading('13.1 NDC Reporting', level=2)
ndc_tests = [
    ('NDC achievement summary', 'Generate report showing cumulative emission reductions toward NDC target'),
    ('ITMO transfer impact', 'Report shows how ITMO transfers affect NDC achievement (corresponding adjustments)'),
    ('Sectoral breakdown', 'NDC report broken down by sector: Energy, Forestry, Agriculture, Waste'),
    ('Time series', 'NDC report shows annual progress from 2020 to current year'),
    ('Gap analysis', 'If NDC target at risk, report highlights gap and mitigation options'),
]
for test, expected in ndc_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('13.2 National Inventory', level=2)
inv_tests = [
    ('Total issued credits', 'Report: total credits issued by year, methodology, province'),
    ('Total retired credits', 'Report: total credits retired by purpose, beneficiary, year'),
    ('Total transferred (ITMOs)', 'Report: total ITMOs transferred by destination country, vintage'),
    ('Active credits', 'Report: credits currently held by each account'),
    ('Buffer pool status', 'Report: buffer pool allocations, releases, draws by project'),
]
for test, expected in inv_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('13.3 Provincial Statistics', level=2)
prov_tests = [
    ('Credits by province', 'Mashonaland West: 5000 tCO2e, Manicaland: 12000 tCO2e, etc.'),
    ('Projects by province', 'Map showing project locations and statuses by province'),
    ('Retirements by province', 'Breakdown of where retired credits originated'),
    ('Verifier coverage by province', 'Which verifiers active in which provinces'),
]
for test, expected in prov_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('13.4 Authorized ITMOs', level=2)
itmo_report_tests = [
    ('Authorization register', 'List of all host country authorizations with UNFCCC IDs'),
    ('Transfer log', 'All international transfers with buyer country, quantity, date'),
    ('Corresponding adjustment log', 'All NDC adjustments with amounts and dates'),
    ('Authorization expiry tracking', 'Alerts for authorizations expiring within 6 months'),
]
for test, expected in itmo_report_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('13.5 Sector Reports', level=2)
sector_tests = [
    ('Forestry sector report', 'All ARR (Afforestation, Reforestation, Revegetation) projects'),
    ('Energy sector report', 'All renewable energy projects: solar, wind, hydro'),
    ('Agriculture sector report', 'All agriculture and agroforestry projects'),
    ('Waste sector report', 'All waste management and methane capture projects'),
    ('Cross-sector comparison', 'Compare emission reductions, credit prices, verification times across sectors'),
]
for test, expected in sector_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 14. Performance & Load Tests
doc.add_heading('14. Performance & Load Tests', level=1)
doc.add_paragraph('Tests to ensure the platform performs under expected and peak loads.')

perf_tests = [
    ('Page load time', 'Dashboard loads in < 3 seconds under normal load'),
    ('Login response time', 'Login API responds in < 500ms'),
    ('Project list load', '100 projects load in < 2 seconds'),
    ('Concurrent users', '50 simultaneous logins → All succeed within 5 seconds'),
    ('Concurrent transactions', '20 simultaneous credit purchases → No data corruption, all recorded'),
    ('Large file upload', 'Upload 100MB evidence ZIP → Completes in < 30 seconds'),
    ('Bulk export', 'Export 10,000 audit events → Completes in < 10 seconds'),
    ('Database query performance', 'Complex report query (joins across 5 tables) → < 2 seconds'),
    ('Memory usage', 'Frontend memory usage < 200MB after 1 hour of use'),
    ('API rate limit', '1000 requests/minute from single IP → Rate limited gracefully'),
    ('Search performance', 'Full-text search across 50,000 records → < 1 second'),
    ('Map rendering', 'GIS map with 6 layers renders in < 3 seconds'),
]
for test, expected in perf_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 15. Blockchain Integration Tests
doc.add_heading('15. Blockchain Integration Tests', level=1)
doc.add_paragraph(
    'Tests for Hyperledger Fabric blockchain integration and immutable record keeping. '
    'The ZAI-CTS blockchain network consists of: 1 Orderer (solo), 4 Peers across 3 organizations '
    '(ZCR, ZiCMA, Market), and the carbon-credit-chaincode smart contract.'
)

doc.add_heading('15.1 Blockchain Network Architecture', level=2)
arch = [
    ('Orderer Node', 'orderer.zai-cts.gov.zw:7050 — Orders transactions into blocks'),
    ('ZCR Peer 0', 'peer0.zcr.zai-cts.gov.zw:7051 — Carbon Registry endorser'),
    ('ZCR Peer 1', 'peer1.zcr.zai-cts.gov.zw:8051 — Carbon Registry endorser (HA)'),
    ('ZiCMA Peer 0', 'peer0.zicma.zai-cts.gov.zw:9051 — Regulatory endorser'),
    ('Market Peer 0', 'peer0.market.zai-cts.gov.zw:10051 — Marketplace endorser'),
    ('Channel', 'zai-cts-channel — Shared ledger for all organizations'),
    ('Chaincode', 'carbon-credit-chaincode — Go smart contract for credit lifecycle'),
]
for component, desc in arch:
    p = doc.add_paragraph()
    p.add_run(f'{component}: ').bold = True
    p.add_run(desc)

doc.add_heading('15.2 Blockchain Setup Procedure', level=2)
setup_steps = [
    '1. Generate crypto materials: docker run hyperledger/fabric-tools cryptogen generate --config=crypto-config.yaml',
    '2. Generate genesis block: configtxgen -profile ZaiCtsOrdererGenesis -outputBlock system-genesis-block/genesis.block',
    '3. Create channel tx: configtxgen -profile ZaiCtsChannel -outputCreateChannelTx channel-artifacts/zai-cts-channel.tx',
    '4. Start containers: docker-compose up -d (starts orderer + 4 peers + CLI)',
    '5. Create channel: peer channel create -c zai-cts-channel -f zai-cts-channel.tx',
    '6. Join all peers: peer channel join (run for each peer with correct env vars)',
    '7. Deploy chaincode: peer lifecycle chaincode install + commit',
    '8. Update backend services to use Fabric SDK (fabric-sdk-py) for real transactions',
]
for step in setup_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('15.3 Functional Blockchain Tests', level=2)
blockchain_tests = [
    ('Credit issuance on-chain', 'Issue 1000 credits → Transaction recorded on Fabric with unique TX ID (fabric:...)'),
    ('Transfer on-chain', 'Transfer credits to buyer → Blockchain updated with new owner, previous owner removed'),
    ('Retirement on-chain', 'Retire credits → Blockchain marked as retired, immutable, certificate hash stored'),
    ('Serial number uniqueness', 'Attempt duplicate serial on blockchain → Rejected by chaincode (serial exists)'),
    ('Chaincode validation', 'Invalid transaction (negative quantity) → Rejected by chaincode validation'),
    ('Block confirmation time', 'Transaction submitted → Confirmed in < 5 seconds (solo orderer)'),
    ('Network partition recovery', 'Temporary network partition → Transactions queued, replayed on recovery'),
    ('Peer synchronization', 'All 4 peer nodes have identical ledger state after each block'),
    ('Endorsement policy', 'Transaction without required ZCR+ZiCMA endorsements → Rejected'),
    ('Private data collection', 'Sensitive buyer data stored in private collection, not on public ledger'),
    ('Chaincode upgrade', 'Upgrade carbon-credit-chaincode → Existing data preserved, new features active'),
    ('Query history', 'Query full transaction history for serial ZWE-RE-2025-00001 → Complete audit trail'),
    ('Cross-organization query', 'ZiCMA queries credits issued by ZCR → Results returned from shared channel'),
    ('Immutable audit', 'Attempt to modify historical block → Rejected by consensus (hash chain broken)'),
]
for test, expected in blockchain_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('15.4 Blockchain Failure Scenarios', level=2)
fail_tests = [
    ('Orderer down', 'Orderer container stops → No new blocks, existing ledger readable, queue transactions'),
    ('Single peer down', 'peer0.zcr fails → Other 3 peers maintain consensus, transactions continue'),
    ('Majority peers down', '3 of 4 peers fail → Network halts until peers recover (safety over liveness)'),
    ('Chaincode panic', 'Bug in chaincode → Container restarts, state preserved, transaction re-evaluated'),
    ('Ledger corruption', 'Corrupted block detected → Peer rejects block, fetches from other peers'),
]
for test, expected in fail_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('15.5 Merkle Root Anchoring Tests (Option B Implementation)', level=2)
doc.add_paragraph(
    'The ZAI-CTS platform implements a pragmatic Merkle root anchoring pattern as an interim ' \
    'blockchain integration (Option B). PostgreSQL remains the source of truth. Each credit entry ' \
    'receives a deterministic SHA-256 hash. Unanchored entries are batched; a Merkle tree is built ' \
    'and its root stored in anchor_batches with a hash chain. A mock fabric_tx_id is recorded; ' \
    'future migration will write the Merkle root to Hyperledger Fabric via the zai-anchor chaincode.'
)

anchor_tests = [
    ('Unanchored count', 'GET /anchors/status/unanchored-count → Returns count of entries not yet anchored'),
    ('Create anchor batch', 'POST /anchors/batch → Batches unanchored entries, computes Merkle root, stores hash chain'),
    ('Merkle root format', 'Merkle root is 64-character lowercase hex SHA-256 string'),
    ('Hash chain continuity', 'Each new anchor references previous_anchor_hash; genesis has null previous'),
    ('Verify anchor', 'GET /anchors/{id}/verify → Recomputes Merkle root from DB entries, compares to stored root'),
    ('Tamper detection', 'Manually edit entry_hash in DB → Verify anchor returns "tampered" (roots mismatch)'),
    ('Reconcile chain', 'GET /anchors/reconcile → Validates every anchor + checks chain continuity across all batches'),
    ('Chain break detection', 'Corrupt previous_anchor_hash → Reconcile flags chain_continuous=false'),
    ('Mock Fabric TX', 'Each anchor stores fabric_tx_id=fabric:<uuid> and fabric_block_number auto-increment'),
    ('Empty batch rejection', 'POST /anchors/batch with 0 unanchored entries → Returns 400 "no_entries"'),
    ('Anchor idempotency', 'Same unanchored entries anchored twice → Second batch only includes newly created entries'),
    ('Frontend anchoring tab', 'Registry console shows "Anchoring" tab with unanchored count, create batch, verify, reconcile'),
    ('Gateway routing', 'GET /anchors/* routes through API Gateway (port 8082) to carbon-registry-service (port 8102)'),
]
for test, expected in anchor_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_heading('15.6 Future Fabric Migration Path', level=2)
doc.add_paragraph(
    'When Hyperledger Fabric is fully deployed: 1) Replace mock fabric_tx_id with real transaction ID. ' \
    '2) Deploy zai-anchor chaincode (Node.js) to zai-cts-channel. 3) Update AnchoringService to call ' \
    'Fabric SDK (createAnchor) after PostgreSQL commit. 4) Use queryAnchor and verifyAnchor chaincode ' \
    'methods for cross-check. 5) Maintain PostgreSQL as primary query store; Fabric as immutable notary.'
)

doc.add_page_break()

# 16. API Contract & Integration Tests
doc.add_heading('16. API Contract & Integration Tests', level=1)
doc.add_paragraph('Tests for API gateway routing, microservice communication, and contract compliance.')

api_tests = [
    ('Gateway routing', 'GET /api/v1/projects → Routed to carbon-registry-service:8102'),
    ('Gateway routing', 'POST /api/v1/auth/login → Routed to carbon-registry-service:8102'),
    ('Gateway routing', 'GET /api/v1/gis/assess → Routed to gis-service:8104'),
    ('Gateway routing', 'GET /api/v1/marketplace/prices → Routed to marketplace-service:8106'),
    ('Service unavailable', 'GIS service down → Gateway returns 503 with retry-after header'),
    ('Timeout handling', 'GIS assessment > 30s → Gateway returns 504, background job continues'),
    ('Request validation', 'POST /api/v1/projects with missing title → 400 Bad Request'),
    ('Authentication required', 'GET /api/v1/admin/users without token → 401 Unauthorized'),
    ('Authorization check', 'Developer token attempts POST /api/v1/credits/issue → 403 Forbidden'),
    ('CORS headers', 'Preflight OPTIONS request → Correct CORS headers returned'),
    ('Content negotiation', 'Accept: application/xml → 406 Not Acceptable (JSON only)'),
    ('API versioning', 'GET /api/v1/projects vs /api/v2/projects → Different response schemas'),
    ('Rate limiting', '100 requests in 1 minute → 429 Too Many Requests'),
    ('OpenAPI validation', 'All endpoints match OpenAPI spec → No contract violations'),
    ('Webhook delivery', 'Credit issued → Webhook POST to subscribed URL within 5 seconds'),
    ('Webhook retry', 'Webhook delivery fails → Retried 3 times with exponential backoff'),
]
for test, expected in api_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 17. Notification & Alert Tests
doc.add_heading('17. Notification & Alert Tests', level=1)
doc.add_paragraph('Tests for email, SMS, and in-app notifications.')

notif_tests = [
    ('Registration confirmation email', 'New user registers → Email sent with confirmation link'),
    ('Approval notification', 'Admin approves user → Email sent: "Your account is approved"'),
    ('Project submission alert', 'Developer submits project → ZiCMA admin receives email alert'),
    ('Verification assignment', 'Verifier assigned to case → Email with case details and deadline'),
    ('Verification complete', 'All verification steps complete → ZiCMA receives notification'),
    ('Credit issuance notification', 'Credits issued → Developer receives email with serial numbers'),
    ('Marketplace sale', 'Credits sold → Seller receives email with transaction details'),
    ('Password reset', 'User requests reset → Email with secure reset link (expires in 1 hour)'),
    ('MFA code', 'User enables MFA → SMS with TOTP setup QR code'),
    ('Login from new device', 'Login from unrecognized IP → Email alert: "New device detected"'),
    ('Failed login attempts', '5 failed attempts → Email: "Suspicious activity detected"'),
    ('System maintenance', 'Scheduled maintenance → Email alert 24 hours in advance'),
    ('In-app alerts', 'New verification case assigned → Red badge on Verification tab'),
    ('Push notifications', 'Browser push enabled → Desktop notification for critical alerts'),
    ('Notification preferences', 'User disables email → Only in-app alerts delivered'),
    ('Bulk notification', 'System alert to all 1000 users → All delivered within 5 minutes'),
]
for test, expected in notif_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 18. Mobile & Responsive Tests
doc.add_heading('18. Mobile & Responsive Tests', level=1)
doc.add_paragraph('Tests for usability across devices and screen sizes.')

mobile_tests = [
    ('Desktop viewport', '1920x1080 → Full layout with sidebar expanded'),
    ('Laptop viewport', '1366x768 → Sidebar collapsible, all features accessible'),
    ('Tablet viewport', '768x1024 → Sidebar becomes hamburger menu, cards stack'),
    ('Mobile viewport', '375x667 → Single column, touch-friendly buttons (> 44px)'),
    ('Mobile login', 'Login form fits on mobile screen without horizontal scroll'),
    ('Mobile project list', 'Project cards stack vertically, readable on mobile'),
    ('Mobile map', 'GIS map supports pinch-to-zoom, pan gestures'),
    ('Mobile file upload', 'Evidence upload works from mobile camera/gallery'),
    ('Touch targets', 'All buttons and links >= 44x44px touch target'),
    ('Font scaling', 'Browser font size 200% → No text truncation or overlap'),
    ('Orientation change', 'Rotate tablet from portrait to landscape → Layout adjusts smoothly'),
    ('Print stylesheet', 'Print project report → Clean print layout without UI chrome'),
    ('Offline indicator', 'Network disconnected → Banner shows "You are offline"'),
    ('Slow network', '3G connection → Loading skeletons shown, progressive loading'),
]
for test, expected in mobile_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 19. Accessibility (WCAG) Tests
doc.add_heading('19. Accessibility (WCAG) Tests', level=1)
doc.add_paragraph('Tests for WCAG 2.1 AA compliance.')

a11y_tests = [
    ('Keyboard navigation', 'Tab through entire login form → All fields and buttons reachable'),
    ('Focus indicators', 'Tab to Login button → Visible focus ring (3px solid)'),
    ('Screen reader labels', 'Email field → Announced as "Email, required, text field"'),
    ('ARIA roles', 'Alert message → Announced with role="alert" and polite/live region'),
    ('Color contrast', 'Text on background → Ratio >= 4.5:1 for normal text, 3:1 for large text'),
    ('Color independence', 'Status badges → Not relying solely on color (icon + text + color)'),
    ('Form errors', 'Submit empty form → Error message linked to field via aria-describedby'),
    ('Skip links', 'Press Tab on page load → "Skip to main content" link appears first'),
    ('Page titles', 'Each tab has unique <title> → "Carbon Registry - ZAI-CTS"'),
    ('Heading hierarchy', 'Page has single H1, logical H2-H6 order → No skipped levels'),
    ('Alt text', 'All images have descriptive alt text → No empty or missing alt'),
    ('PDF accessibility', 'Generated reports are tagged PDFs → Screen reader readable'),
    ('Zoom 200%', 'Browser zoom to 200% → No horizontal scroll, all content visible'),
    ('Motion preferences', 'prefers-reduced-motion: reduce → Animations disabled'),
    ('Cognitive accessibility', 'Complex workflows have step indicators and progress bars'),
]
for test, expected in a11y_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 20. Disaster Recovery & Backup Tests
doc.add_heading('20. Disaster Recovery & Backup Tests', level=1)
doc.add_paragraph('Tests for business continuity and data recovery.')

dr_tests = [
    ('Database backup', 'Nightly PostgreSQL backup → Verified restorable, < 1% data loss'),
    ('Point-in-time recovery', 'Restore database to 4 hours ago → All transactions recovered'),
    ('Blockchain backup', 'Peer node failure → Remaining 3 peers maintain consensus, new peer syncs'),
    ('Frontend deployment', 'Deploy new frontend version → Zero-downtime with rolling deployment'),
    ('Service restart', 'Kill and restart carbon-registry-service → < 30 second downtime'),
    ('Gateway failover', 'Primary gateway fails → Traffic routes to secondary gateway'),
    ('Database failover', 'Primary PostgreSQL fails → Automatic failover to replica < 10 seconds'),
    ('Data corruption detection', 'Corrupted audit_event row detected → Alert sent, data quarantined'),
    ('Ransomware simulation', 'Encrypted files detected → Isolate system, restore from clean backup'),
    ('Load balancer health', 'Unhealthy backend removed from rotation → Traffic only to healthy nodes'),
    ('Configuration backup', 'Gateway config changed → Previous config backed up, rollback possible'),
    ('Log retention', 'Audit logs retained for 7 years → Searchable, tamper-evident'),
]
for test, expected in dr_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 21. Session & Authentication Security
doc.add_heading('21. Session & Authentication Security', level=1)
doc.add_paragraph('Tests for session management, token handling, and authentication security.')

auth_tests = [
    ('Token expiry', 'JWT expires after 1 hour → User redirected to login'),
    ('Token refresh', 'Refresh token rotates on use → Old refresh token invalidated'),
    ('Concurrent sessions', 'User logs in from 2 browsers → Both sessions valid, user sees active sessions'),
    ('Session termination', 'User clicks "Logout" → Token blacklisted, cannot be reused'),
    ('Admin force logout', 'Admin terminates user session → User immediately logged out'),
    ('Password complexity', 'Password "123456" → Rejected. Requires: 12+ chars, upper, lower, number, special'),
    ('Password history', 'User changes to previous password → Rejected. Last 12 passwords blocked'),
    ('Account lockout', '10 failed logins → Account locked for 30 minutes'),
    ('MFA enforcement', 'Admin enables MFA for all → Users without MFA cannot access'),
    ('MFA bypass attempt', 'Attacker intercepts TOTP → Cannot reuse (time-based, 30-second window)'),
    ('OAuth integration', 'Login with Microsoft Entra ID → Account linked, SSO works'),
    ('SAML integration', 'Login with government SSO → SAML assertion validated, roles mapped'),
    ('Certificate pinning', 'API client rejects certificate from untrusted CA → Connection refused'),
    ('Credential stuffing', '1000 common passwords tried → Rate limited, IP blocked after 50 attempts'),
]
for test, expected in auth_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 22. Data Integrity & Migration Tests
doc.add_heading('22. Data Integrity & Migration Tests', level=1)
doc.add_paragraph('Tests for data consistency, schema migrations, and data transformation.')

migrate_tests = [
    ('Schema migration forward', 'Apply migration 0003 → New column added, existing data preserved'),
    ('Schema migration rollback', 'Rollback migration 0003 → Column removed, data backed up'),
    ('Data seeding', 'Fresh database → Seed data: admin user, default organization, test projects'),
    ('Cross-service consistency', 'Project deleted in registry → Related credits, evidence, audits also removed'),
    ('Foreign key integrity', 'Delete organization with projects → CASCADE or RESTRICT as defined'),
    ('Data type validation', 'Insert text into numeric field → Database rejects with type error'),
    ('Unicode support', 'Project title in Shona: "Kushandura KweMhene" → Stored and displayed correctly'),
    ('Date timezone handling', 'Event at 2026-06-22 14:00 Harare → Stored as UTC, displayed in local time'),
    ('Large dataset import', 'Import 10,000 historical projects → All imported, validation errors logged'),
    ('CSV export/import roundtrip', 'Export projects to CSV → Re-import → Data matches original'),
    ('Audit log completeness', 'Every CREATE/UPDATE/DELETE logged → No missing audit events'),
    ('Checksum validation', 'Evidence file checksum stored → Re-verify on access → Match confirmed'),
]
for test, expected in migrate_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 23. Search & Export Functionality
doc.add_heading('23. Search & Export Functionality', level=1)
doc.add_paragraph('Tests for search, filtering, and data export capabilities.')

search_tests = [
    ('Full-text search', 'Search "Kariba" → Returns Kariba projects, evidence, audit events'),
    ('Fuzzy search', 'Search "Karbia" (typo) → Suggests "Kariba" with did-you-mean'),
    ('Filter by status', 'Filter projects by "approved" → Only approved projects shown'),
    ('Filter by date range', 'Filter audit events: 2026-01-01 to 2026-06-30 → Events in range only'),
    ('Filter by province', 'Filter by "Mashonaland West" → Only projects in that province'),
    ('Combined filters', 'Status: approved + Province: Manicaland + Methodology: VM0047 → Matching projects'),
    ('Sort by date', 'Sort projects by created_at descending → Newest first'),
    ('Sort by tCO2e', 'Sort projects by estimated_annual_tco2e → Largest to smallest'),
    ('Pagination', '100 projects, page size 10 → 10 pages, correct navigation'),
    ('Export to CSV', 'Export filtered project list → CSV with all visible columns'),
    ('Export to Excel', 'Export marketplace listings → .xlsx with formatting'),
    ('Export to PDF', 'Export project report → PDF with logo, header, footer, page numbers'),
    ('Export to GeoJSON', 'Export project boundaries → GeoJSON for GIS analysis'),
    ('Scheduled export', 'Schedule monthly report → Auto-generated and emailed on 1st of month'),
    ('Export large dataset', 'Export 50,000 audit events → Streamed download, no memory overflow'),
]
for test, expected in search_tests:
    p = doc.add_paragraph()
    p.add_run(f'{test}: ').bold = True
    p.add_run(expected)

doc.add_page_break()

# 24. Troubleshooting
doc.add_heading('24. Troubleshooting', level=1)
doc.add_heading('5.1 Service Health Checks', level=2)
doc.add_paragraph('Run these commands to verify all services are running:')
doc.add_paragraph('curl -s http://127.0.0.1:8082/health   # API Gateway', style='List Bullet')
doc.add_paragraph('curl -s http://127.0.0.1:8102/health   # Carbon Registry', style='List Bullet')
doc.add_paragraph('curl -s http://127.0.0.1:8103/health   # AI Validation', style='List Bullet')
doc.add_paragraph('curl -s http://127.0.0.1:8104/health   # GIS Service', style='List Bullet')
doc.add_paragraph('curl -s http://127.0.0.1:8106/health   # Marketplace', style='List Bullet')
doc.add_paragraph('curl -s http://127.0.0.1:8107/health   # Compliance', style='List Bullet')

doc.add_heading('5.2 Common Issues', level=2)
issues = [
    ('Page stuck on loading', 'Next.js dev server compiling. Wait 30–60 seconds. Refresh page.'),
    ('Login fails with 401', 'Account not yet approved by admin. Contact ZiCMA Administrator.'),
    ('Gateway unhealthy', 'One or more backend services down. Check service health endpoints.'),
    ('GIS layers not loading', 'GIS service (port 8104) may be down. Restart: cd backend/services/gis-service && uvicorn app.main:app --host 127.0.0.1 --port 8104'),
    ('Marketplace prices not updating', 'Marketplace service (port 8106) may need restart.'),
    ('Cannot issue credits', 'Project must be in approved status. Complete verification pipeline first.'),
]
for issue, fix in issues:
    p = doc.add_paragraph()
    p.add_run(f'{issue}: ').bold = True
    p.add_run(fix)

doc.add_heading('5.3 Reset / Clear Cache', level=2)
doc.add_paragraph('If frontend behaves unexpectedly:')
doc.add_paragraph('cd frontend/web-portal && rm -rf .next && npm run dev', style='List Bullet')
doc.add_paragraph('If gateway config changed, restart gateway:', style='List Bullet')
doc.add_paragraph('cd api-gateway && PORT=8082 npx tsx src/server.ts', style='List Bullet')

doc.add_page_break()

# Summary
doc.add_heading('Summary', level=1)
doc.add_paragraph(
    'This guide covers the complete end-to-end testing of the ZAI-CTS platform across all user roles. '
    'The platform supports the full carbon credit lifecycle from project registration through '
    'verification, issuance, marketplace trading, and retirement. Each role has distinct '
    'responsibilities and permissions, enforced through RBAC. All actions are auditable with '
    'immutable event logging.'
)

# Save
doc.save('ZAI-CTS_E2E_Testing_Guide.docx')
print('DOCX created: ZAI-CTS_E2E_Testing_Guide.docx')
